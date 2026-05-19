import os

def load_env_file():
    env_path = os.path.join(os.path.dirname(__file__), '.env')
    if os.path.exists(env_path):
        try:
            with open(env_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    if '=' in line:
                        key, val = line.split('=', 1)
                        key = key.strip()
                        val = val.strip().strip("'").strip('"')
                        os.environ[key] = val
            print("[Env Config] Loaded environment variables from .env file.")
        except Exception as e:
            print(f"[Env Config Warning] Failed to read .env file: {e}")

load_env_file()

from flask import Flask, render_template, redirect, url_for, flash, request, jsonify, session, send_file
import pandas as pd
import io
import json
from datetime import datetime, timedelta
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import threading
import time
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', str(s))]

def sanitize_filename(name):
    """Removes invalid filename characters and truncates to safe length."""
    # Keep only alphanumeric, spaces, underscores, and dashes
    s = str(name)
    s = "".join([c for c in s if c.isalnum() or c in (' ', '_', '-')]).strip()
    return s[:100] if s else "Download"

# Global state for background scheduler
scheduler_status = {
    'running': False,
    'progress': 0,
    'message': 'No active generation.',
    'last_result': None,
    'cancel_requested': False
}
from models import db, User, Teacher, Classroom, Section, Subject, Setting, Schedule, ScheduleRun, ActivityLog

def log_activity(actor_username, role, action, module):
    try:
        log = ActivityLog(
            actor_username=actor_username,
            role=role,
            action=action,
            module=module,
            timestamp=datetime.now()
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        print(f"Error logging activity: {e}")

app = Flask(__name__)
# Load configuration
app.config['SECRET_KEY'] = 'your-secret-key-here'  # Change this in production
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'database.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)

@app.before_request
def check_session_timeout():
    if current_user.is_authenticated:
        now = datetime.now()
        last_active_str = session.get('last_activity')
        
        if last_active_str:
            try:
                last_active = datetime.fromisoformat(last_active_str)
                if now - last_active > timedelta(minutes=15):
                    log_activity(current_user.username, current_user.role, 'Session Expired Logout', 'Authentication')
                    logout_user()
                    session.pop('last_activity', None)
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'error': 'Session expired', 'redirect': url_for('login')}), 401
                    flash('Session Expired. Please log in again.', 'warning')
                    return redirect(url_for('login'))
            except ValueError:
                pass
        
        session['last_activity'] = now.isoformat()
# --- AUTO MIGRATION ---
with app.app_context():
    try:
        from sqlalchemy import inspect, text
        inspector = inspect(db.engine)
        # Ensure new tables like ScheduleRun are created
        db.create_all()

        columns = [c['name'] for c in inspector.get_columns('section')]
        if 'is_section_a' not in columns:
            db.session.execute(text("ALTER TABLE section ADD COLUMN is_section_a BOOLEAN DEFAULT 0"))
            db.session.commit()
            print("Auto-migration: Added is_section_a to section table.")
            
        schedule_cols = [c['name'] for c in inspector.get_columns('schedule')]
        if schedule_cols and 'run_id' not in schedule_cols:
            db.session.execute(text("ALTER TABLE schedule ADD COLUMN run_id INTEGER REFERENCES schedule_run(id)"))
            db.session.commit()
            print("Auto-migration: Added run_id to schedule table.")

        if schedule_cols and 'is_soft_break_override' not in schedule_cols:
            db.session.execute(text("ALTER TABLE schedule ADD COLUMN is_soft_break_override BOOLEAN DEFAULT 0"))
            db.session.commit()
            print("Auto-migration: Added is_soft_break_override to schedule table.")

            
        user_columns = [c['name'] for c in inspector.get_columns('user')]
        if 'is_active' not in user_columns:
            db.session.execute(text("ALTER TABLE user ADD COLUMN is_active BOOLEAN DEFAULT 1"))
            db.session.commit()
            print("Auto-migration: Added is_active to user table.")

        if 'password_updated_at' not in user_columns:
            db.session.execute(text("ALTER TABLE user ADD COLUMN password_updated_at DATETIME"))
            db.session.commit()
            print("Auto-migration: Added password_updated_at to user table.")

        if 'is_super_admin' not in user_columns:
            db.session.execute(text("ALTER TABLE user ADD COLUMN is_super_admin BOOLEAN DEFAULT 0"))
            db.session.commit()
            print("Auto-migration: Added is_super_admin to user table.")

        if 'failed_login_attempts' not in user_columns:
            db.session.execute(text("ALTER TABLE user ADD COLUMN failed_login_attempts INTEGER DEFAULT 0"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN locked_until DATETIME"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN security_question VARCHAR(200)"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN security_answer VARCHAR(200)"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN recovery_otp VARCHAR(6)"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN recovery_otp_expiry DATETIME"))
            db.session.execute(text("ALTER TABLE user ADD COLUMN recovery_email VARCHAR(120)"))
            db.session.commit()
            print("Auto-migration: Added security & recovery fields to user table.")

        # Ensure at least one admin is promoted to Super Admin if there are admins but none are Super Admin
        super_admin_exists = User.query.filter_by(role='admin', is_super_admin=True).first() is not None
        if not super_admin_exists:
            first_admin = User.query.filter_by(role='admin').order_by(User.id.asc()).first()
            if first_admin:
                first_admin.is_super_admin = True
                db.session.commit()
                print(f"Auto-migration: Promoted first existing admin '{first_admin.username}' to Super Admin.")
                try:
                    log = ActivityLog(
                        actor_username='system / setup / initial admin creator',
                        role='system',
                        action='Promoted First Admin to Super Admin Account (Migration)',
                        module='Users',
                        timestamp=datetime.now()
                    )
                    db.session.add(log)
                    db.session.commit()
                except Exception as e:
                    print(f"Error logging migration activity: {e}")

        teacher_columns = [c['name'] for c in inspector.get_columns('teacher')]
        if 'is_hybrid' not in teacher_columns:
            db.session.execute(text("ALTER TABLE teacher ADD COLUMN is_hybrid BOOLEAN DEFAULT 0"))
            db.session.commit()
            print("Auto-migration: Added is_hybrid to teacher table.")

        if 'stay_window_hours' not in teacher_columns:
            db.session.execute(text("ALTER TABLE teacher ADD COLUMN stay_window_hours INTEGER DEFAULT 9"))
            db.session.commit()
            print("Auto-migration: Added stay_window_hours to teacher table.")
            
        # --- ORPHAN CLEANUP ON STARTUP ---
        def cleanup_orphans():
            # Teachers
            teacher_users = User.query.filter_by(role='teacher').all()
            for u in teacher_users:
                if not Teacher.query.get(u.related_id):
                    db.session.delete(u)
            # Sections
            student_users = User.query.filter_by(role='student').all()
            for u in student_users:
                if not Section.query.get(u.related_id):
                    db.session.delete(u)
            db.session.commit()
        
        cleanup_orphans()
        print("Startup: Orphaned accounts cleaned up.")
    except Exception as e:
        print(f"Auto-migration failed: {e}")

login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    user = User.query.get(int(user_id))
    if user and user.is_active:
        return user
    return None

@app.context_processor
def inject_school_name():
    setting = Setting.query.filter_by(key='school_name').first()
    sy_setting = Setting.query.filter_by(key='school_year').first()
    return {
        'school_name': setting.value if setting else 'Andres M. Luciano High School',
        'school_year': sy_setting.value if sy_setting else '',
        'is_complete': is_complete
    }

def time_to_min(t_str):
    h, m = map(int, t_str.split(':'))
    return h * 60 + m

def min_to_time(m):
    return f"{m//60:02d}:{m%60:02d}"

def get_default_password():
    try:
        setting = Setting.query.filter_by(key='default_password').first()
        if setting and setting.value:
            return setting.value
    except:
        pass
    return '123456'

def normalize_gl(gl):
    if not gl: return gl
    gl_s = str(gl).strip()
    # Idempotent check: if already "Grade X", return normalized version
    if re.match(r'^Grade\s+\d+$', gl_s, flags=re.IGNORECASE):
        num = re.search(r'\d+', gl_s).group()
        return f"Grade {num}"
    
    # Remove any existing "Grade", "Gr", or "G" prefix or spaces
    clean = re.sub(r'^(Grade|Gr|G)\s*', '', gl_s, flags=re.IGNORECASE)
    
    if clean.isdigit():
        return f"Grade {clean}"
    return clean

import traceback

@app.errorhandler(500)
def handle_500(error):
    with open('crash_log.txt', 'a') as f:
        f.write(f"\n--- {datetime.now()} ---\n")
        f.write(traceback.format_exc())
        f.write("-" * 20 + "\n")
    return "Internal Server Error. Diagnostic info saved to crash_log.txt", 500

def is_complete(obj):
    try:
        if isinstance(obj, Teacher):
            return all([obj.name, obj.department, obj.grade_levels, obj.subjects, obj.max_hours_per_day, obj.stay_window_hours])
        elif isinstance(obj, Classroom):
            return all([obj.name, obj.room_type, obj.building])
        elif isinstance(obj, Section):
            base = all([obj.name, obj.department, obj.grade_level, obj.adviser_id, obj.room_id])
            if obj.department == 'SHS':
                return base and bool(obj.track)
            return base
        elif isinstance(obj, Subject):
            return all([obj.name, obj.department, obj.duration_mins, obj.meetings_per_week, obj.grade_level])
    except:
        return False
    return False

# --- ROUTES ---

@app.route('/')
def index():
    if current_user.is_authenticated:
        if current_user.role == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif current_user.role == 'teacher':
            return redirect(url_for('teacher_dashboard'))
        elif current_user.role == 'student':
            return redirect(url_for('student_dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        role = request.form.get('role')
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter((User.username == username) | (User.recovery_email == username)).first()
        if not user:
            flash("Account not found.")
            return redirect(url_for('login'))
            
        if not user.is_active:
            flash("This account is suspended or inactive.")
            return redirect(url_for('login'))
        if not is_debug_mode() and user.is_super_admin and user.locked_until and user.locked_until > datetime.now():
            session['recovery_username'] = user.username
            session['show_recovery_link'] = True
            flash("Account is temporarily locked due to multiple failed login attempts.", "error")
            return redirect(url_for('login'))
            
        pass_ok = check_password_hash(user.password, password)
        role_ok = (user.role == role)
        
        if not pass_ok or not role_ok:
            if user.is_super_admin:
                if is_debug_mode():
                    log_activity(user.username, user.role, 'Failed Login Attempt (Development Bypass Active)', 'Authentication')
                    flash("Incorrect credentials. (Lockout is bypassed in Development Mode)", "error")
                else:
                    user.failed_login_attempts += 1
                    attempts_left = 3 - user.failed_login_attempts
                    if attempts_left <= 0:
                        user.locked_until = datetime.now() + timedelta(hours=24)
                        db.session.commit()
                        log_activity(user.username, user.role, 'Account Locked (Multiple Failed Logins)', 'Authentication')
                        session['recovery_username'] = user.username
                        session['show_recovery_link'] = True
                        flash("Account is temporarily locked due to multiple failed login attempts.", "error")
                    else:
                        db.session.commit()
                        log_activity(user.username, user.role, f'Failed Login Attempt ({user.failed_login_attempts}/3)', 'Authentication')
                        flash(f"Incorrect credentials. Super Admin account will lock in {attempts_left} more failed attempt(s).", "error")
                return redirect(url_for('login'))
            else:
                if not pass_ok and not role_ok:
                    flash("Invalid role and incorrect password.", "error")
                elif not pass_ok:
                    flash("Incorrect password. Please try again.", "error")
                elif not role_ok:
                    flash("Invalid role selected.", "error")
                return redirect(url_for('login'))
        
        # Successful login: Reset attempts if super admin
        if user.is_super_admin:
            user.failed_login_attempts = 0
            user.locked_until = None
            db.session.commit()
        # Check if linked record still exists (Orphan Detection)
        exists = True
        if role == 'teacher':
            exists = Teacher.query.get(user.related_id) is not None
        elif role == 'student':
            exists = Section.query.get(user.related_id) is not None
        
        if not exists:
            db.session.delete(user)
            db.session.commit()
            flash("Account not found or has been removed.")
            return redirect(url_for('login'))
            
        login_user(user)
        session.permanent = True
        session['last_activity'] = datetime.now().isoformat()
        log_activity(user.username, user.role, 'User Logged In', 'Authentication')
        if role == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif role == 'teacher':
            return redirect(url_for('teacher_dashboard'))
        elif role == 'student':
            return redirect(url_for('student_dashboard'))
    return render_template('login.html')

def is_debug_mode():
    # Explicit environment standardization using ENV = "development" | "production"
    # Secure by default: If ENV is not explicitly set to 'development', it defaults to 'production'.
    env_mode = os.environ.get('ENV', 'production').strip().lower()
    return env_mode == 'development'

def send_otp_email(recipient_email, otp):
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    
    smtp_server = os.environ.get('SMTP_SERVER', 'smtp.gmail.com')
    try:
        smtp_port = int(os.environ.get('SMTP_PORT', '587'))
    except ValueError:
        smtp_port = 587
        
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_password = os.environ.get('SMTP_PASSWORD', '').replace(' ', '')
    
    is_configured = (
        smtp_user and smtp_password and 
        smtp_user.strip() != '' and smtp_password.strip() != '' and 
        'your_gmail' not in smtp_user and 'your_gmail' not in smtp_password
    )
    
    if not is_configured:
        if is_debug_mode():
            print("\n" + "="*80)
            print(f"[DEVELOPMENT MOCK EMAIL]")
            print(f"To: {recipient_email}")
            print(f"Subject: Super Admin Account Recovery - Verification OTP")
            print(f"Body: Your 6-digit Verification OTP code is {otp}")
            print("="*80 + "\n")
            session['smtp_mode'] = 'development'
            return True
        else:
            print("[SMTP Error] Production Mode is active, but SMTP credentials are not configured in environment.")
            session['smtp_mode'] = 'production'
            return False
        
    session['smtp_mode'] = 'production'
    
    msg = MIMEMultipart()
    msg['From'] = smtp_user if smtp_user else "no-reply@amlhs.edu.ph"
    msg['To'] = recipient_email
    msg['Subject'] = "Super Admin Account Recovery - Verification OTP"
    
    body = f"""Hello,

You have requested password recovery for the Super Admin account of the Class Scheduling System.

Your 6-digit Verification OTP code is:

{otp}

This code is valid for 10 minutes. If you did not request this code, please secure your account immediately.

Sincerely,
Class Scheduling System Admin Service
"""
    msg.attach(MIMEText(body, 'plain'))
    
    try:
        if smtp_port == 465:
            server = smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=10)
        else:
            server = smtplib.SMTP(smtp_server, smtp_port, timeout=10)
            server.starttls()
            
        server.login(smtp_user, smtp_password)
        server.send_message(msg)
        server.quit()
        print(f"[SMTP Success] OTP email sent successfully to {recipient_email}.")
        return True
    except Exception as e:
        if is_debug_mode():
            print(f"[SMTP Error] Failed to send OTP email to {recipient_email}: {e}")
        else:
            print(f"[SMTP Error] Production email dispatch failed. Ensure App Password is correct.")
        return False

@app.route('/recover_account', methods=['GET', 'POST'])
def recover_account():
    username = session.get('recovery_username')
    step = 1
    
    if request.method == 'POST':
        action = request.form.get('action')
        form_username = request.form.get('username', '').strip()
        
        if action == 'step1':
            form_email = request.form.get('recovery_email', '').strip()
            user = User.query.filter_by(is_super_admin=True).first()
            
            # Use 'admin@amlhs.edu.ph' as fallback if DB field is empty
            expected_email = user.recovery_email if user and user.recovery_email else "admin@amlhs.edu.ph"
            
            if not user or form_email.lower() != expected_email.lower():
                flash("Incorrect registered email.", "error")
                return render_template('recover.html', step=1, smtp_error=False)
                
            session['recovery_username'] = user.username
            
            import random
            otp = f"{random.randint(100000, 999999)}"
            user.recovery_otp = otp
            user.recovery_otp_expiry = datetime.now() + timedelta(minutes=10)
            db.session.commit()
            
            # Real email sending
            if send_otp_email(expected_email, otp):
                if session.get('smtp_mode') == 'development':
                    flash("Development Mode: SMTP credentials are not configured in the .env file. The OTP has been printed to the server terminal console.", "info")
                else:
                    flash("A 6-digit OTP code has been securely sent to your registered email address.", "success")
                log_activity(user.username, user.role, 'Requested Account Recovery OTP', 'Authentication')
                return render_template('recover.html', step=2, username=user.username, smtp_mode=session.get('smtp_mode'), dev_otp=otp)
            else:
                if not is_debug_mode():
                    flash("Failed to send verification email. In Production Mode, a valid SMTP configuration (SMTP_USER/SMTP_PASSWORD) is strictly required.", "error")
                else:
                    flash("Failed to send verification email. Please verify SMTP_USER and SMTP_PASSWORD settings in the .env file and ensure TLS/SSL is allowed.", "error")
                return render_template('recover.html', step=1, smtp_error=True)
            
        elif action == 'step2':
            user = User.query.filter_by(username=form_username, is_super_admin=True).first()
            otp = request.form.get('otp', '').strip()
            
            if is_debug_mode() or (user and user.recovery_otp and user.recovery_otp == otp and user.recovery_otp_expiry > datetime.now()):
                session['recovery_otp_verified'] = True
                return render_template('recover.html', step=3, username=form_username)
            else:
                flash("Invalid or expired OTP.", "error")
                return render_template('recover.html', step=2, username=form_username, smtp_mode=session.get('smtp_mode'), dev_otp=user.recovery_otp if user else "")
                
        elif action == 'step3':
            if not session.get('recovery_otp_verified'):
                flash("Unauthorized access.", "error")
                return redirect(url_for('login'))
                
            user = User.query.filter_by(username=form_username, is_super_admin=True).first()
            new_password = request.form.get('new_password', '').strip()
            
            import re
            if len(new_password) < 8 or not re.search('[a-zA-Z]', new_password) or not re.search('[0-9]', new_password):
                flash('Password must be at least 8 characters and contain both letters and numbers.', 'error')
                return render_template('recover.html', step=3, username=form_username)
                
            user.password = generate_password_hash(new_password)
            user.password_updated_at = datetime.now()
            user.locked_until = None
            user.failed_login_attempts = 0
            user.recovery_otp = None
            user.recovery_otp_expiry = None
            db.session.commit()
            
            log_activity(user.username, user.role, 'Account Recovered & Password Reset via OTP', 'Authentication')
            session.pop('show_recovery_link', None)
            session.pop('recovery_otp_verified', None)
            flash("Your account has been successfully unlocked and your password has been reset. Please log in.", "success")
            return redirect(url_for('login'))
 
    return render_template('recover.html', step=1, username=username, smtp_error=False)

@app.route('/logout')
@login_required
def logout():
    log_activity(current_user.username, current_user.role, 'User Logged Out', 'Authentication')
    logout_user()
    session.pop('last_activity', None)
    return redirect(url_for('login'))

@app.route('/auto_logout')
def auto_logout():
    if current_user.is_authenticated:
        log_activity(current_user.username, current_user.role, 'Session Expired Logout', 'Authentication')
        logout_user()
        session.pop('last_activity', None)
        flash('Session Expired. Please log in again.', 'warning')
    return redirect(url_for('login'))

@app.route('/admin')
@login_required
def admin_dashboard():
    if current_user.role != 'admin':
        return redirect(url_for('index'))
    # Fetch counts for the dashboard stats cards
    teachers = Teacher.query.all()
    sections = Section.query.all()
    classrooms = Classroom.query.all()
    subjects = Subject.query.all()
    return render_template('admin_dashboard.html', 
        teachers=teachers, sections=sections, 
        classrooms=classrooms, subjects=subjects)



@app.route('/admin/update_recovery_email', methods=['POST'])
@login_required
def update_recovery_email():
    if current_user.role != 'admin' or not current_user.is_super_admin:
        flash("Unauthorized access.", "error")
        return redirect(url_for('admin_dashboard'))
        
    new_email = request.form.get('recovery_email', '').strip()
    
    if not new_email:
        flash("Email cannot be empty.", "error")
        return redirect(url_for('admin_dashboard'))
        
    current_user.recovery_email = new_email
    db.session.commit()
    
    log_activity(current_user.username, current_user.role, 'Updated Registered Recovery Email', 'Authentication')
    flash("Recovery email updated successfully.", "success")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/teachers', methods=['GET', 'POST'])
@login_required
def admin_teachers():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            department = request.form.get('department')
            
            # Duplicate prevention
            existing = Teacher.query.filter(Teacher.name.ilike(name)).first()
            if existing:
                error_msg = f'Error: Teacher "{name}" is already registered.'
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'error': error_msg}), 400
                flash(error_msg, 'error')
                return redirect(url_for('admin_teachers'))
                
            grade_levels = ','.join([normalize_gl(g) for g in request.form.getlist('grade_levels')])
            max_hours = int(request.form.get('max_hours_per_day', 6))
            stay_window = int(request.form.get('stay_window_hours', 9))
            is_master = True if request.form.get('is_master') else False
            handle_sec_a = True if request.form.get('handle_sec_a') else False
            preferred_days = request.form.get('preferred_days', 'Mon-Fri')
            subjects = ','.join(request.form.getlist('subjects'))
            
            teacher = Teacher(
                name=name, department=department, grade_levels=grade_levels,
                max_hours_per_day=max_hours, stay_window_hours=stay_window,
                is_master=is_master,
                handle_sec_a=handle_sec_a, preferred_days=preferred_days, subjects=subjects
            )
            db.session.add(teacher)
            db.session.commit()
            
            # Auto-create User account for the teacher
            if not User.query.filter_by(username=name).first():
                user = User(username=name, password=generate_password_hash(get_default_password()), role='teacher', related_id=teacher.id)
                db.session.add(user)
                db.session.commit()
            log_activity(current_user.username, current_user.role, f"Created Teacher Profile: {name}", "Teachers")
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': True, 'msg': 'Teacher added!', 'row_html': render_template('partials/_teacher_row.html', t=teacher, teacher_loads={}, is_complete=is_complete), 'dept': department})
                
            flash(f'Teacher added successfully. Default password is {get_default_password()}')
            return redirect(url_for('admin_teachers'))
        except Exception as e:
            db.session.rollback()
            error_msg = f"Error adding teacher: {str(e)}"
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 500
            flash(error_msg, 'error')
            return redirect(url_for('admin_teachers'))
    
    # Ensure we have the latest data from the DB for recalculation
    db.session.expire_all()
    
    # Optimized Calculation of loads for "Undertime" detection
    teacher_loads = {}
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    if active_run:
        from sqlalchemy import func
        # Sum duration_mins from Subject entries linked via Schedule
        results = db.session.query(
            Schedule.teacher_id, 
            func.sum(Subject.duration_mins)
        ).join(Subject).filter(
            Schedule.run_id == active_run.id
        ).group_by(Schedule.teacher_id).all()
        
        for t_id, total_mins in results:
            if t_id:
                teacher_loads[t_id] = round(total_mins / 60, 1)
    
    # The loads are already in hours and rounded from the query results loop above

    # Calculate school days from settings
    active_days_setting = Setting.query.filter_by(key='active_days').first()
    school_days = len(active_days_setting.value.split(',')) if active_days_setting and active_days_setting.value else 5

    # Standard Weekly Load from settings
    standard_weekly_load_setting = Setting.query.filter_by(key='standard_weekly_load').first()
    standard_weekly_load = float(standard_weekly_load_setting.value) if standard_weekly_load_setting and standard_weekly_load_setting.value else 30.0

    # Sort teachers by department (JHS, Both, SHS order or similar) and name
    teachers = [t for t in Teacher.query.all() if t.name != 'TBA']
    teachers.sort(key=lambda x: (x.department, natural_sort_key(x.name)))
    subjects_list = Subject.query.all()
    subjects_list.sort(key=lambda x: (x.department, natural_sort_key(x.name)))
    return render_template('admin_teachers.html', 
                          teachers=teachers, 
                          subjects=subjects_list, 
                          teacher_loads=teacher_loads, 
                          school_days=school_days,
                          standard_weekly_load=standard_weekly_load)

@app.route('/admin/import/<module>/preview', methods=['POST'])
@login_required
def admin_import_preview(module):
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    
    text_data = request.form.get('text_data')
    file = request.files.get('file')
    raw_rows = []

    # --- SMART MAPPING CONFIGURATION ---
    smart_mappings = {
        'teachers': {
            'name': ['name', 'full name', 'teacher name', 'teacher', 'instructor'],
            'department': ['department', 'dept', 'level'],
            'max_hours_per_day': ['teaching load', 'max hours', 'max hours per day', 'hours', 'limit'],
            'stay_window_hours': ['stay window', 'stay hours', 'school stay', 'stay hours per day'],
            'grade_levels': ['grade levels', 'grades', 'handled grades'],
            'is_master': ['master teacher', 'master', 'is master', 'master?'],
            'preferred_days': ['preferred days', 'days', 'schedule preference'],
            'subjects': ['subjects', 'assigned subjects']
        },
        'classrooms': {
            'name': ['room name', 'name', 'room', 'rm', 'rm name', 'classroom'],
            'room_type': ['room type', 'type', 'category'],
            'building': ['building', 'bldg', 'location', 'phase']
        },
        'sections': {
            'name': ['section name', 'name', 'section', 'sec name'],
            'department': ['dept', 'department', 'level'],
            'grade_level': ['grade', 'grade level', 'level', 'gr'],
            'track': ['track', 'strand'],
            'adviser_id': ['adviser', 'adviser name', 'teacher', 'assigned teacher'],
            'room_id': ['room', 'room name', 'assigned room', 'assigned home room', 'home room'],
            'is_section_a': ['is section a', 'section a', 'handle a', 'priority', 'is a?']
        },
        'subjects': {
            'name': ['subject name', 'name', 'subject', 'subj'],
            'department': ['dept', 'department', 'level'],
            'duration_mins': ['duration', 'mins', 'duration mins', 'time', 'length'],
            'meetings_per_week': ['freq', 'frequency', 'meetings per week', 'sessions'],
            'requires_lab': ['lab', 'requires lab', 'requires laboratory', 'is lab', 'needs lab', 'laboratory?'],
            'grade_level': ['grade', 'grade level', 'level', 'gr', 'grades'],
            'track': ['track', 'strand']
        }
    }

    def get_field_match(header, module_map):
        if pd.isna(header): return None
        # Clean header: lowercase and remove spaces, underscores, hyphens, and common punctuation
        def clean(s):
            return str(s).lower().strip().replace(' ', '').replace('_', '').replace('-', '').replace('?', '').replace('.', '')
        
        h_clean = clean(header)
        # Direct match check
        for field, aliases in module_map.items():
            if h_clean == clean(field):
                return field
            for alias in aliases:
                if h_clean == clean(alias):
                    return field
        return None

    if file and file.filename != '':
        try:
            if file.filename.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                # Explicit engine for modern Excel files
                engine = 'openpyxl' if file.filename.endswith('.xlsx') else None
                dfs = pd.read_excel(file, sheet_name=None, engine=engine)
                df = pd.concat(dfs.values(), ignore_index=True)
            
            df = df.dropna(how='all')
            module_map = smart_mappings.get(module, {})
            
            # Map columns to internal keys
            col_to_field = {}
            for col in df.columns:
                matched_field = get_field_match(col, module_map)
                if matched_field:
                    col_to_field[col] = matched_field
            
            final_data = []
            for _, row in df.iterrows():
                mapped_row = {field: row[col] for col, field in col_to_field.items() if not pd.isna(row[col])}
                if str(mapped_row.get('name', '')).strip() in ['Sample Teacher', 'Sample Room', 'Sample Section', 'Sample Subject']:
                    continue
                final_data.append(mapped_row)
            raw_rows = final_data
        except Exception as e:
            return jsonify({'error': f'File parsing error: {str(e)}'}), 400
    elif text_data:
        # Automatic Separator Detection
        lines = [l.strip() for l in text_data.strip().split('\n') if l.strip()]
        if not lines: return jsonify({'data': []})
        
        # Check first line for best separator
        sample = lines[0]
        separator = '|'
        for s in ['\t', '|', ',', ';']:
            if s in sample:
                separator = s
                break

        for line in lines:
            parts = [p.strip() for p in line.split(separator)]
            if module == 'teachers':
                row = {
                    'name': parts[0] if len(parts) > 0 and parts[0] else '',
                    'department': parts[1] if len(parts) > 1 and parts[1] else 'JHS',
                    'max_hours_per_day': parts[2] if len(parts) > 2 and parts[2] else '6',
                    'stay_window_hours': parts[3] if len(parts) > 3 and parts[3] else '9',
                    'grade_levels': parts[4] if len(parts) > 4 and parts[4] else '',
                    'is_master': parts[5] if len(parts) > 5 and parts[5] else 'No',
                    'preferred_days': parts[6] if len(parts) > 6 and parts[6] else 'Mon-Fri',
                    'subjects': parts[7] if len(parts) > 7 and parts[7] else ''
                }
                if row['name']: raw_rows.append(row)
            elif module == 'classrooms':
                row = {
                    'name': parts[0] if len(parts) > 0 and parts[0] else '',
                    'room_type': parts[1] if len(parts) > 1 and parts[1] else 'Room',
                    'building': parts[2] if len(parts) > 2 and parts[2] else ''
                }
                if row['name']: raw_rows.append(row)
            elif module == 'sections':
                row = {
                    'name': parts[0] if len(parts) > 0 and parts[0] else '',
                    'department': parts[1] if len(parts) > 1 and parts[1] else 'JHS',
                    'grade_level': parts[2] if len(parts) > 2 and parts[2] else '',
                    'track': parts[3] if len(parts) > 3 and parts[3] else '',
                    'adviser_id': parts[4] if len(parts) > 4 and parts[4] else '',
                    'room_id': parts[5] if len(parts) > 5 and parts[5] else '',
                    'is_section_a': parts[6] if len(parts) > 6 and parts[6] else 'No'
                }
                if row['name']: raw_rows.append(row)
            elif module == 'subjects':
                row = {
                    'name': parts[0] if len(parts) > 0 and parts[0] else '',
                    'department': parts[1] if len(parts) > 1 and parts[1] else 'JHS',
                    'duration_mins': parts[2] if len(parts) > 2 and parts[2] else '60',
                    'meetings_per_week': parts[3] if len(parts) > 3 and parts[3] else '4',
                    'requires_lab': parts[4] if len(parts) > 4 and parts[4] else 'No',
                    'grade_level': parts[5] if len(parts) > 5 and parts[5] else '',
                    'track': parts[6] if len(parts) > 6 and parts[6] else ''
                }
                if row['name']: raw_rows.append(row)

    # Boolean and Numeric Normalization
    def to_bool(val):
        if pd.isna(val) or val is None: return False
        s = str(val).lower().strip()
        return s in ['yes', 'true', '1', 'y', 't']

    for row in raw_rows:
        row['_errors'] = []
        
        # Module-Specific Defaults and Validation
        if module == 'teachers':
            dept = str(row.get('department', '')).upper().strip()
            if 'BOTH' in dept or ('JHS' in dept and 'SHS' in dept): row['department'] = 'Both'
            elif 'SHS' in dept: row['department'] = 'SHS'
            else: row['department'] = 'JHS'

            if not row.get('name'):
                row['_errors'].append('Missing Teacher Name')
            else:
                existing = Teacher.query.filter(Teacher.name.ilike(str(row['name']).strip())).first()
                if existing:
                    row['_errors'].append(f'Duplicate record: Teacher "{row["name"]}" is already registered')

        elif module == 'classrooms':
            if not row.get('room_type'): row['room_type'] = 'Room'
            # Building Normalization
            bldg = str(row.get('building', '')).strip().upper()
            if any(x in bldg for x in ['BOTH', 'JHS,SHS', 'JHS/SHS', 'JHS & SHS']):
                row['building'] = 'Both'
            elif 'SHS' in bldg: row['building'] = 'SHS'
            elif 'JHS' in bldg: row['building'] = 'JHS'
            
            if not row.get('building'): row['_errors'].append('Missing Building')
            if not row.get('name'):
                row['_errors'].append('Missing Room Name')
            else:
                existing = Classroom.query.filter(Classroom.name.ilike(str(row['name']).strip())).first()
                if existing:
                    row['_errors'].append(f'Duplicate record: Room "{row["name"]}" already exists')
            
        elif module == 'sections':
            dept = str(row.get('department', '')).upper().strip()
            if 'SHS' in dept: row['department'] = 'SHS'
            else: row['department'] = 'JHS'

            if not row.get('name'): row['_errors'].append('Missing Section Name')
            if not row.get('department'): row['_errors'].append('Missing Department')
            if not row.get('grade_level'): row['_errors'].append('Missing Grade Level')
            
            # Duplicate prevention (including grade suffix check)
            if row.get('name') and row.get('grade_level'):
                s_name = str(row['name']).strip()
                s_gl = str(row['grade_level']).strip().upper().replace('GRADE', '').strip()
                
                existing = Section.query.filter(Section.name.ilike(s_name)).first()
                if existing:
                    row['_errors'].append(f'Duplicate record: Section "{s_name}" already exists')

            # Adviser and Room name lookup for preview
            a_name = row.get('adviser_id')
            r_name = row.get('room_id')
            row['_adviser_name'] = a_name
            row['_room_name'] = r_name
            
            if a_name:
                adviser = Teacher.query.filter_by(name=str(a_name).strip()).first()
                if not adviser: row['_errors'].append(f'Adviser "{a_name}" not found')
            if r_name:
                room = Classroom.query.filter_by(name=str(r_name).strip()).first()
                if not room: row['_errors'].append(f'Room "{r_name}" not found')

        elif module == 'subjects':
            dept = str(row.get('department', '')).upper().strip()
            if 'BOTH' in dept or ('JHS' in dept and 'SHS' in dept): row['department'] = 'Both'
            elif 'SHS' in dept: row['department'] = 'SHS'
            else: row['department'] = 'JHS'

            if not row.get('name'): row['_errors'].append('Missing Subject Name')
            if not row.get('department'): row['_errors'].append('Missing Department')
            
            if row.get('name') and row.get('department'):
                s_name = str(row['name']).strip()
                s_dept = str(row['department']).strip().upper()
                existing = Subject.query.filter(Subject.name.ilike(s_name), Subject.department == s_dept).first()
                if existing:
                    row['_errors'].append(f'Duplicate record: Subject "{s_name}" already exists in {s_dept}')
            
            # Defaults
            if not row.get('requires_lab'): row['requires_lab'] = 'No'
            if not row.get('duration_mins'): row['duration_mins'] = 60
            if not row.get('meetings_per_week'): row['meetings_per_week'] = 1
            
            # Grade Level Parsing
            gl = str(row.get('grade_level', '')).strip().upper()
            if 'ALL JHS' in gl: row['grade_level'] = '7,8,9,10'
            elif 'ALL SHS' in gl: row['grade_level'] = '11,12'
            elif 'GRADE' in gl:
                row['grade_level'] = gl.replace('GRADE', '').strip()
            
            # Track Parsing (Multiple values)
            track = str(row.get('track', '')).strip()
            if track:
                # Standardize separators
                track = track.replace('/', ',').replace(';', ',')
                row['track'] = ', '.join([t.strip() for t in track.split(',') if t.strip()])

        # Boolean and Numeric Normalization
        for bool_field in ['is_master', 'requires_lab', 'is_section_a', 'handle_sec_a']:
            if bool_field in row:
                val = row[bool_field]
                if pd.isna(val) or val is None or str(val).strip() == '':
                    row[bool_field] = False
                else:
                    s = str(val).lower().strip()
                    if s in ['yes', 'true', '1', 'y', 't']:
                        row[bool_field] = True
                    elif s in ['no', 'false', '0', 'n', 'f']:
                        row[bool_field] = False
                    else:
                        row['_errors'].append(f'Invalid value for {bool_field}: "{val}". Must be Yes or No.')
                        row[bool_field] = False
        
        # Numeric checks
        for num_field in ['max_hours_per_day', 'stay_window_hours', 'duration_mins', 'meetings_per_week']:
            if num_field in row and row[num_field]:
                try: 
                    row[num_field] = int(float(row[num_field]))
                    if num_field == 'meetings_per_week' and row[num_field] > 5:
                        row['_errors'].append(f'Warning: {row[num_field]} meetings/week exceeds recommended limit (5)')
                except: 
                    pass
        
        # Special case for Grade Level (keep as string if it's a range/list)
        if 'grade_level' in row and row['grade_level']:
            gl_val = str(row['grade_level']).strip()
            if gl_val.isdigit():
                row['grade_level'] = int(gl_val)
            else:
                row['grade_level'] = gl_val

        # Final Completeness Check
        row['_is_complete'] = len(row['_errors']) == 0

    return jsonify({'data': raw_rows})

@app.route('/admin/bulk_delete/<module>', methods=['POST'])
@login_required
def admin_bulk_delete(module):
    if current_user.role != 'admin':
        if request.is_json:
            return jsonify({'error': 'Unauthorized'}), 403
        return redirect(url_for('index'))
    
    # Support both JSON (AJAX) and Form Data (Standard Submit)
    if request.is_json:
        selected_ids = request.json.get('selected_ids', [])
    else:
        selected_ids = request.form.getlist('selected_ids')

    if not selected_ids:
        if request.is_json:
            return jsonify({'error': 'No records selected'}), 400
        flash("No records selected.")
        return redirect(url_for(f'admin_{module}'))
        
    success_ids = []
    errors = []
    
    try:
        for item_id in selected_ids:
            try:
                # Module-Specific Dependency Checks
                if module == 'teachers':
                    item = Teacher.query.get(item_id)
                    # Cascade remove from Sections (Adviser)
                    Section.query.filter_by(adviser_id=item_id).update({'adviser_id': None})
                    # Cascade remove schedules
                    Schedule.query.filter_by(teacher_id=item_id).delete()
                    # Cascade: Hard Delete User account
                    User.query.filter_by(role='teacher', related_id=item_id).delete()
                    db.session.delete(item)
                elif module == 'classrooms':
                    item = Classroom.query.get(item_id)
                    # Cascade remove from Sections (Home Room)
                    Section.query.filter_by(room_id=item_id).update({'room_id': None})
                    # Cascade remove schedules
                    Schedule.query.filter_by(room_id=item_id).delete()
                    db.session.delete(item)
                elif module == 'sections':
                    item = Section.query.get(item_id)
                    # Cascade remove schedules
                    Schedule.query.filter_by(section_id=item_id).delete()
                    # Cascade: Hard Delete User account
                    User.query.filter_by(role='student', related_id=item_id).delete()
                    db.session.delete(item)
                elif module == 'subjects':
                    item = Subject.query.get(item_id)
                    # Cascade remove schedules
                    Schedule.query.filter_by(subject_id=item_id).delete()
                    # Check if assigned to Teachers
                    all_teachers = Teacher.query.all()
                    assigned_teachers = [t.name for t in all_teachers if item.name in [s.strip() for s in str(t.subjects or '').split(',')]]
                    if assigned_teachers:
                        errors.append(f"Subject '{item.name}' is assigned to Teachers: {', '.join(assigned_teachers[:2])}...")
                        continue
                    db.session.delete(item)
                success_ids.append(item_id)
            except Exception as e:
                errors.append(f"Error deleting ID {item_id}: {str(e)}")
        
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        if request.is_json:
            return jsonify({'error': f"Database error: {str(e)}"}), 500
        flash(f"Database error: {str(e)}", "error")
        return redirect(url_for(f'admin_{module}'))
        
    if request.is_json:
        return jsonify({
            'success': True,
            'success_count': len(success_ids),
            'errors': errors
        })
    
    if success_ids:
        flash(f"Successfully deleted {len(success_ids)} items.", "success")
    if errors:
        for err in errors: flash(err, "error")
    return redirect(url_for(f'admin_{module}'))

@app.route('/admin/import/<module>/confirm', methods=['POST'])
@login_required
def admin_import_confirm(module):
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json.get('data', [])
    count = 0
    try:
        for row in data:
            if module == 'teachers':
                name = str(row.get('name', '')).strip()
                if Teacher.query.filter(Teacher.name.ilike(name)).first():
                    continue # Skip duplicates
                obj = Teacher(
                    name=name,
                    department=row.get('department', 'JHS'),
                    grade_levels=','.join([normalize_gl(g.strip()) for g in str(row.get('grade_levels', '')).split(',') if g.strip()]),
                    max_hours_per_day=int(row.get('max_hours_per_day', 6)),
                    stay_window_hours=int(row.get('stay_window_hours', 9)),
                    is_master=bool(row.get('is_master', False)),
                    preferred_days=row.get('preferred_days', 'Mon-Fri'),
                    subjects=str(row.get('subjects', ''))
                )
            elif module == 'classrooms':
                name = str(row.get('name', '')).strip()
                if Classroom.query.filter(Classroom.name.ilike(name)).first():
                    continue # Skip duplicates
                obj = Classroom(
                    name=name,
                    room_type=row.get('room_type', 'Lecture'),
                    building=row.get('building', '')
                )
            elif module == 'sections':
                name = str(row.get('name', '')).strip()
                grade_level = normalize_gl(row.get('grade_level', '7'))
                
                if Section.query.filter(Section.name.ilike(name)).first():
                    continue # Skip duplicates
                
                # Adviser and Room mapping
                adviser_name = row.get('adviser_id')
                room_name = row.get('room_id')
                adviser = Teacher.query.filter_by(name=adviser_name).first() if adviser_name else None
                room = Classroom.query.filter_by(name=room_name).first() if room_name else None
                
                obj = Section(
                    name=name,
                    department=row.get('department', 'JHS'),
                    grade_level=grade_level,
                    track=row.get('track', ''),
                    adviser_id=adviser.id if adviser else None,
                    room_id=room.id if room else None,
                    is_section_a=bool(row.get('is_section_a', False))
                )
            elif module == 'subjects':
                name = str(row.get('name', '')).strip()
                department = str(row.get('department', 'JHS')).strip().upper()
                if Subject.query.filter(Subject.name.ilike(name), Subject.department == department).first():
                    continue # Skip duplicates
                
                gl = normalize_gl(row.get('grade_level') or '7')
                
                obj = Subject(
                    name=name,
                    department=department,
                    duration_mins=int(row.get('duration_mins') or 60),
                    meetings_per_week=int(row.get('meetings_per_week') or 1),
                    requires_lab=bool(row.get('requires_lab', False)),
                    grade_level=gl,
                    track=row.get('track', '')
                )
            
            db.session.add(obj)
            db.session.flush() # Get ID for user account
            
            if module == 'teachers':
                if not User.query.filter_by(username=obj.name).first():
                    user = User(username=obj.name, password=generate_password_hash(get_default_password()), role='teacher', related_id=obj.id)
                    db.session.add(user)
            elif module == 'sections':
                if not User.query.filter_by(username=obj.name).first():
                    user = User(username=obj.name, password=generate_password_hash(get_default_password()), role='student', related_id=obj.id)
                    db.session.add(user)
            
            count += 1
        
        db.session.commit()
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/admin/teachers/delete/<int:id>')
@login_required
def delete_teacher(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    # Cascade remove from Sections (Adviser)
    Section.query.filter_by(adviser_id=id).update({'adviser_id': None})
    # Cascade remove schedules
    Schedule.query.filter_by(teacher_id=id).delete()
    # Cascade: Hard Delete User account
    User.query.filter_by(role='teacher', related_id=id).delete()
    t = Teacher.query.get_or_404(id)
    t_name = t.name
    db.session.delete(t)
    db.session.commit()
    log_activity(current_user.username, current_user.role, f"Deleted Teacher Profile: {t_name}", "Teachers")
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': 'Teacher deleted.'})
    flash('Teacher deleted.')
    return redirect(url_for('admin_teachers'))

@app.route('/admin/teachers/edit/<int:id>', methods=['POST'])
@login_required
def edit_teacher(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    try:
        teacher = Teacher.query.get_or_404(id)
        old_name = teacher.name
        name = request.form.get('name', '').strip()
        
        # Duplicate prevention (excluding self)
        existing = Teacher.query.filter(Teacher.name.ilike(name), Teacher.id != id).first()
        if existing:
            error_msg = f'Error: A teacher with the name "{name}" is already registered.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_teachers'))

        teacher.name = name
        teacher.department = request.form.get('department')
        teacher.grade_levels = ','.join([normalize_gl(g) for g in request.form.getlist('grade_levels')])
        teacher.max_hours_per_day = int(request.form.get('max_hours_per_day', 6))
        teacher.stay_window_hours = int(request.form.get('stay_window_hours', 9))
        teacher.is_master = True if request.form.get('is_master') else False
        teacher.handle_sec_a = True if request.form.get('handle_sec_a') else False
        teacher.preferred_days = request.form.get('preferred_days', 'Mon-Fri')
        teacher.subjects = ','.join(request.form.getlist('subjects'))
        
        # Sync User account
        user = User.query.filter_by(username=old_name, role='teacher', related_id=teacher.id).first()
        if user:
            user.username = teacher.name
            
        db.session.commit()
        log_activity(current_user.username, current_user.role, f"Edited Teacher Profile: {teacher.name}", "Teachers")
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                'success': True, 
                'message': 'Teacher updated.',
                'teacher': {
                    'id': teacher.id,
                    'name': teacher.name,
                    'department': teacher.department,
                    'grade_levels': teacher.grade_levels,
                    'max_hours': teacher.max_hours_per_day,
                    'stay_window': teacher.stay_window_hours,
                    'is_master': teacher.is_master,
                    'handle_sec_a': teacher.handle_sec_a,
                    'subjects': teacher.subjects,
                    'is_active': teacher.is_active
                }
            })
            
        flash('Teacher updated.')
        return redirect(url_for('admin_teachers'))
    except Exception as e:
        db.session.rollback()
        error_msg = f'Error updating teacher: {str(e)}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': False, 'error': error_msg}), 500
        flash(error_msg, 'error')
        return redirect(url_for('admin_teachers'))

@app.route('/admin/classrooms', methods=['GET', 'POST'])
@login_required
def admin_classrooms():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            
            # Duplicate prevention
            existing = Classroom.query.filter(Classroom.name.ilike(name)).first()
            if existing:
                error_msg = f'Error: Room "{name}" already exists.'
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'error': error_msg}), 400
                flash(error_msg, 'error')
                return redirect(url_for('admin_classrooms'))
    
            room_type = request.form.get('room_type')
            building = request.form.get('building')
            room = Classroom(name=name, room_type=room_type, building=building)
            db.session.add(room)
            db.session.commit()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': True, 'msg': 'Classroom added!', 'row_html': render_template('partials/_classroom_row.html', c=room, is_complete=is_complete), 'bldg': building})
                
            flash('Classroom added.')
            return redirect(url_for('admin_classrooms'))
        except Exception as e:
            db.session.rollback()
            error_msg = f"Error adding classroom: {str(e)}"
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 500
            flash(error_msg, 'error')
            return redirect(url_for('admin_classrooms'))
    # Sort classrooms by department, building, name
    classrooms = Classroom.query.all()
    classrooms.sort(key=lambda x: (x.building, natural_sort_key(x.name)))
    return render_template('admin_classrooms.html', classrooms=classrooms)

@app.route('/admin/classrooms/delete/<int:id>')
@login_required
def delete_classroom(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    # Cascade remove from Sections (Home Room)
    Section.query.filter_by(room_id=id).update({'room_id': None})
    # Cascade remove schedules
    Schedule.query.filter_by(room_id=id).delete()
    db.session.delete(Classroom.query.get_or_404(id))
    db.session.commit()
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': 'Classroom deleted.'})
    flash('Classroom deleted.')
    return redirect(url_for('admin_classrooms'))

@app.route('/admin/classrooms/edit/<int:id>', methods=['POST'])
@login_required
def edit_classroom(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    try:
        room = Classroom.query.get_or_404(id)
        name = request.form.get('name', '').strip()
        
        # Duplicate prevention (excluding self)
        existing = Classroom.query.filter(Classroom.name.ilike(name), Classroom.id != id).first()
        if existing:
            error_msg = f'Error: Room "{name}" already exists.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_classrooms'))

        room.name = name
        room.room_type = request.form.get('room_type')
        room.building = request.form.get('building')
        db.session.commit()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': True, 'msg': 'Classroom updated!'})
            
        flash('Classroom updated.')
        return redirect(url_for('admin_classrooms'))
    except Exception as e:
        db.session.rollback()
        error_msg = f'Error updating room: {str(e)}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': False, 'error': error_msg}), 500
        flash(error_msg, 'error')
        return redirect(url_for('admin_classrooms'))

@app.route('/admin/sections', methods=['GET', 'POST'])
@login_required
def admin_sections():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        name = request.form.get('name')
        department = request.form.get('department')
        grade_level = request.form.get('grade_level')
        # Combined Track logic
        track = request.form.get('track_select')
        
        adviser_id = request.form.get('adviser_id')
        room_id = request.form.get('room_id')
        
        is_section_a = True if request.form.get('is_section_a') else False
        
        # Robust ID parsing
        adviser_id = int(adviser_id) if (adviser_id and adviser_id.strip() and adviser_id.lower() != 'none') else None
        room_id = int(room_id) if (room_id and room_id.strip() and room_id.lower() != 'none') else None

        # Enforce naming convention: Name_Grade
        if f"_{grade_level}" not in name:
            name = f"{name}_{grade_level}"
            
        # Duplicate prevention (Case-insensitive name check)
        existing = Section.query.filter(Section.name.ilike(name)).first()
        if existing:
            error_msg = f'Error: Section "{name}" already exists.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_sections'))

        try:
            # Shift validation for room assignment
            if room_id:
                settings_dict = {s.key: s.value for s in Setting.query.all()}
                inc_shift = get_section_shift(department, grade_level, settings_dict)
                is_valid, v_msg = validate_room_assignment(int(room_id), None, inc_shift, settings_dict)
                if not is_valid:
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'success': False, 'error': v_msg}), 400
                    flash(v_msg, 'error')
                    return redirect(url_for('admin_sections'))

            section = Section(name=name, department=department, grade_level=grade_level, track=track, adviser_id=adviser_id, room_id=room_id, is_section_a=is_section_a)
            db.session.add(section)
            db.session.commit()
            
            # Auto-create User account for the section
            if not User.query.filter_by(username=name).first():
                user = User(username=name, password=generate_password_hash(get_default_password()), role='student', related_id=section.id)
                db.session.add(user)
                db.session.commit()
                
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                # Deterministic check for row partial
                has_all = all([getattr(section, f, None) not in [None, ''] for f in ['name', 'department', 'grade_level', 'room_id']])
                return jsonify({
                    'success': True, 
                    'msg': 'Section added!', 
                    'row_html': render_template('partials/_section_row.html', s=section, is_complete=lambda x: has_all), 
                    'dept': department
                })
                
        except Exception as e:
            db.session.rollback()
            err_msg = f"Error adding section: {str(e)}"
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': err_msg}), 500
            flash(err_msg, 'error')
            return redirect(url_for('admin_sections'))
            
        # Extract custom track
        if track:
            track_setting = Setting.query.filter_by(key='available_tracks').first()
            if not track_setting:
                track_setting = Setting(key='available_tracks', value='TVL,STEM,HUMSS')
                db.session.add(track_setting)
            current_tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
            if track and track not in current_tracks:
                current_tracks.append(track)
                track_setting.value = ','.join(sorted(current_tracks))
            db.session.commit()
            
        flash(f'Section added. Default password is {get_default_password()}')
        return redirect(url_for('admin_sections'))
        
    sections = Section.query.all()
    teachers = Teacher.query.all()
    classrooms = Classroom.query.all()
    # Sort sections by department (JHS first), grade level, and name
    sections = Section.query.all()
    sections.sort(key=lambda x: (x.department, natural_sort_key(x.grade_level), natural_sort_key(x.name)))
    # Fetch tracks with proactive initialization
    track_setting = Setting.query.filter_by(key='available_tracks').first()
    if not track_setting:
        # Seed from existing data + defaults
        existing_tracks = set([s.track for s in sections if s.track])
        existing_tracks.update(set([sub.track for sub in Subject.query.all() if sub.track]))
        existing_tracks.update({'TVL', 'STEM', 'HUMSS'})
        track_setting = Setting(key='available_tracks', value=','.join(sorted(list(existing_tracks))))
        db.session.add(track_setting)
        db.session.commit()
        
    tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
    return render_template('admin_sections.html', sections=sections, teachers=teachers, classrooms=classrooms, tracks=tracks)

@app.route('/admin/sections/delete/<int:id>')
@login_required
def delete_section(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    # Cascade remove schedules
    Schedule.query.filter_by(section_id=id).delete()
    # Cascade: Hard Delete User account
    User.query.filter_by(role='student', related_id=id).delete()
    db.session.delete(Section.query.get_or_404(id))
    db.session.commit()
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': 'Section deleted.'})
    flash('Section deleted.')
    return redirect(url_for('admin_sections'))

@app.route('/admin/tracks/add', methods=['POST'])
@login_required
def admin_track_add():
    if current_user.role != 'admin': return redirect(url_for('index'))
    new_track = request.form.get('new_track', '').strip()
    
    if not new_track:
        flash('Track name cannot be empty.', 'error')
        return redirect(url_for('admin_sections'))
        
    track_setting = Setting.query.filter_by(key='available_tracks').first()
    if not track_setting:
        track_setting = Setting(key='available_tracks', value='')
        db.session.add(track_setting)
        
    tracks = [t.strip() for t in track_setting.value.split(',') if t.strip()]
    if new_track not in tracks:
        tracks.append(new_track)
        track_setting.value = ','.join(sorted(tracks))
        db.session.commit()
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': f'Track "{new_track}" added!', 'tracks': tracks})
        
    return redirect(url_for('admin_sections'))

@app.route('/admin/tracks/edit', methods=['POST'])
@login_required
def admin_track_edit():
    if current_user.role != 'admin': return redirect(url_for('index'))
    old_track = request.form.get('old_track', '').strip()
    new_track = request.form.get('new_track', '').strip()
    
    if not new_track:
        flash('New track name cannot be empty.')
        return redirect(url_for('admin_sections'))
        
    # Sync all sections and subjects
    Section.query.filter_by(track=old_track).update({Section.track: new_track})
    Subject.query.filter_by(track=old_track).update({Subject.track: new_track})
    
    # Update available_tracks setting
    track_setting = Setting.query.filter_by(key='available_tracks').first()
    if track_setting:
        tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
        if old_track in tracks:
            tracks = [new_track if t == old_track else t for t in tracks]
            track_setting.value = ','.join(sorted(list(set(tracks))))
    else:
        # If setting was missing, create it now with the change
        tracks = set([s.track for s in Section.query.all() if s.track])
        tracks.update({'TVL', 'STEM', 'HUMSS'})
        tracks.add(new_track)
        track_setting = Setting(key='available_tracks', value=','.join(sorted(list(tracks))))
        db.session.add(track_setting)
    
    db.session.commit()
    
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': f'Track "{old_track}" renamed to "{new_track}"', 'tracks': tracks})
        
    flash(f'Track "{old_track}" renamed to "{new_track}". All related records updated.')
    return redirect(url_for('admin_sections'))

@app.route('/admin/tracks/delete/<path:name>')
@login_required
def admin_track_delete(name):
    name = name.strip()
    if current_user.role != 'admin': return redirect(url_for('index'))
    
    # Update all sections and subjects to None
    Section.query.filter_by(track=name).update({Section.track: None})
    Subject.query.filter_by(track=name).update({Subject.track: None})
    
    # Update available_tracks setting
    track_setting = Setting.query.filter_by(key='available_tracks').first()
    if track_setting:
        tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
        if name in tracks:
            tracks = [t for t in tracks if t != name]
            track_setting.value = ','.join(tracks)
            
    db.session.commit()

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()] if track_setting else []
        return jsonify({'success': True, 'msg': f'Track "{name}" deleted.', 'tracks': tracks})

    flash(f'Track "{name}" deleted. All related records cleared.')
    return redirect(url_for('admin_sections'))

@app.route('/admin/sections/edit/<int:id>', methods=['POST'])
@login_required
def edit_section(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    try:
        section = Section.query.get_or_404(id)
        old_name = section.name
        
        # Debug logging to identify missing fields in form submission
        print(f"DEBUG EDIT SECTION: id={id}, form_keys={list(request.form.keys())}")
        
        name = request.form.get('name', '').strip()
        new_grade = request.form.get('grade_level')
        
        # Clean naming: remove existing grade suffix if transitioning (e.g. Diamond_10 -> Diamond_9)
        import re
        name = re.sub(r'_\d+$', '', name) 
        if f"_{new_grade}" not in name:
            name = f"{name}_{new_grade}"
            
        # Duplicate prevention (excluding self) - Case-insensitive name safety
        existing = Section.query.filter(Section.name.ilike(name), Section.id != int(id)).first()
        if existing:
            error_msg = f'Error: Section name must be unique. "{name}" already exists.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_sections'))

        new_track = request.form.get('track_select')
        
        section.name = name
        
        # Force preservation of non-nullable fields
        received_dept = request.form.get('department')
        if received_dept and received_dept.strip():
            section.department = received_dept
        else:
            print(f"DEBUG: Department missing in form, keeping existing: {section.department}", flush=True)
        
        if new_grade:
            section.grade_level = new_grade
        else:
            print(f"DEBUG: Grade level missing in form, keeping existing: {section.grade_level}", flush=True)
            
        section.track = new_track
        
        adviser_id = request.form.get('adviser_id')
        room_id = request.form.get('room_id')
        
        # Robust ID parsing
        section.adviser_id = int(adviser_id) if (adviser_id and adviser_id.strip() and adviser_id.lower() != 'none') else None
        
        parsed_room_id = int(room_id) if (room_id and room_id.strip() and room_id.lower() != 'none') else None
        
        # Shift validation for room assignment
        if parsed_room_id:
            settings_dict = {s.key: s.value for s in Setting.query.all()}
            inc_shift = get_section_shift(section.department, section.grade_level, settings_dict)
            is_valid, v_msg = validate_room_assignment(parsed_room_id, section.id, inc_shift, settings_dict)
            if not is_valid:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'error': v_msg}), 400
                flash(v_msg, 'error')
                return redirect(url_for('admin_sections'))
                
        section.room_id = parsed_room_id
        section.is_section_a = True if request.form.get('is_section_a') else False
    
        # Sync User account
        user = User.query.filter_by(username=old_name, role='student', related_id=section.id).first()
        if user:
            user.username = section.name
            
        db.session.commit()
        
        # Ensure track exists in global settings if used
        if section.track:
            track_setting = Setting.query.filter_by(key='available_tracks').first()
            if not track_setting:
                track_setting = Setting(key='available_tracks', value='TVL,STEM,HUMSS')
                db.session.add(track_setting)
            current_tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
            if section.track and section.track not in current_tracks:
                current_tracks.append(section.track)
                track_setting.value = ','.join(sorted(current_tracks))
                db.session.commit()

        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            # Pre-calculate completion status to avoid runtime template errors
            is_comp_val = all([getattr(section, f, None) not in [None, ''] for f in ['name', 'department', 'grade_level', 'room_id']])
            return jsonify({
                'success': True, 
                'msg': 'Section updated!', 
                'row_html': render_template('partials/_section_row.html', s=section, is_complete=lambda x: is_comp_val)
            })
            
        flash('Section updated.')
        return redirect(url_for('admin_sections'))
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in edit_section: {str(e)}")
        err_msg = f"Error updating section: {str(e)}"
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': False, 'error': err_msg}), 500
        flash(err_msg, 'error')
        return redirect(url_for('admin_sections'))

@app.route('/admin/subjects', methods=['GET', 'POST'])
@login_required
def admin_subjects():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        department = request.form.get('department')
        
        # Duplicate subject prevention (Case-insensitive)
        existing = Subject.query.filter(Subject.name.ilike(name), Subject.department == department).first()
        if existing:
            error_msg = f'Error: Subject "{name}" already exists in {department}.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_subjects'))

        try:
            # Safe parsing with defaults
            duration_raw = request.form.get('duration_mins', '60')
            duration = int(duration_raw) if (duration_raw and duration_raw.isdigit()) else 60
            
            meetings_raw = request.form.get('meetings_per_week', '5')
            meetings = int(meetings_raw) if (meetings_raw and meetings_raw.isdigit()) else 5
            
            track = request.form.get('track_select')
            
            subject = Subject(
                name=name, 
                department=department,
                requires_lab=True if request.form.get('requires_lab') else False,
                duration_mins=duration,
                meetings_per_week=meetings,
                grade_level=normalize_gl(request.form.get('grade_level', '7')),
                track=track
            )
            db.session.add(subject)
            db.session.commit()
            
            # Extract custom track
            if track:
                track_setting = Setting.query.filter_by(key='available_tracks').first()
                if not track_setting:
                    track_setting = Setting(key='available_tracks', value='TVL,STEM,HUMSS')
                    db.session.add(track_setting)
                current_tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
                if track and track not in current_tracks:
                    current_tracks.append(track)
                    track_setting.value = ','.join(sorted(current_tracks))
                db.session.commit()

            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                is_comp_val = all([getattr(subject, f, None) not in [None, ''] for f in ['name', 'department', 'duration_mins', 'grade_level']])
                return jsonify({
                    'success': True, 
                    'msg': 'Subject added!', 
                    'row_html': render_template('partials/_subject_row.html', s=subject, is_complete=lambda x: is_comp_val), 
                    'dept': department
                })

            flash('Subject added.')
            return redirect(url_for('admin_subjects'))
        except Exception as e:
            db.session.rollback()
            print(f"ERROR in admin_subjects POST: {str(e)}") # Console debug
            error_msg = f'Error adding subject: {str(e)}'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 500
            flash(error_msg, 'error')
            return redirect(url_for('admin_subjects'))
        
    # Sort subjects by department, grade level, and name (hiding system subjects)
    subjects = Subject.query.filter_by(is_system=False).all()
    subjects.sort(key=lambda x: (x.department, natural_sort_key(str(x.grade_level)), natural_sort_key(x.name)))
    track_setting = Setting.query.filter_by(key='available_tracks').first()
    if not track_setting:
        # Proactive sync
        all_sects = Section.query.all()
        existing_tracks = set([s.track for s in all_sects if s.track])
        existing_tracks.update(set([sub.track for sub in subjects if sub.track]))
        existing_tracks.update({'TVL', 'STEM', 'HUMSS'})
        track_setting = Setting(key='available_tracks', value=','.join(sorted(list(existing_tracks))))
        db.session.add(track_setting)
        db.session.commit()
        
    tracks = [t.strip() for t in (track_setting.value or '').split(',') if t.strip()]
    return render_template('admin_subjects.html', subjects=subjects, tracks=tracks)

@app.route('/admin/subjects/delete/<int:id>')
@login_required
def delete_subject(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    # Cascade remove schedules
    Schedule.query.filter_by(subject_id=id).delete()
    db.session.delete(Subject.query.get_or_404(id))
    db.session.commit()
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': 'Subject deleted.'})
    flash('Subject deleted.')
    return redirect(url_for('admin_subjects'))

@app.route('/admin/subjects/edit/<int:id>', methods=['POST'])
@login_required
def edit_subject(id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    try:
        subject = Subject.query.get_or_404(id)
        name = request.form.get('name', '').strip()
        department = request.form.get('department')
        
        # Duplicate subject prevention (Case-insensitive, excluding self)
        existing = Subject.query.filter(Subject.name.ilike(name), Subject.department == department, Subject.id != int(id)).first()
        if existing:
            error_msg = f'Error: Subject "{name}" already exists in {department}.'
            print(f"COLLISION: Subject '{name}' (ID: {id}) collided with existing ID {existing.id}")
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('admin_subjects'))
            
        # Debug logging
        print(f"DEBUG EDIT SUBJECT: id={id}, form_keys={list(request.form.keys())}")
        
        # Safe parsing with defaults
        duration_raw = request.form.get('duration_mins')
        subject.duration_mins = int(duration_raw) if (duration_raw and duration_raw.isdigit()) else subject.duration_mins
        
        meetings_raw = request.form.get('meetings_per_week')
        subject.meetings_per_week = int(meetings_raw) if (meetings_raw and meetings_raw.isdigit()) else subject.meetings_per_week
        
        subject.name = name
        
        # Force preservation
        received_dept = request.form.get('department')
        if received_dept and received_dept.strip():
            subject.department = received_dept
        else:
            print(f"DEBUG: Subject department missing, keeping existing: {subject.department}", flush=True)
            
        subject.requires_lab = True if request.form.get('requires_lab') else False
        
        received_gl = request.form.get('grade_level')
        if received_gl:
            subject.grade_level = normalize_gl(received_gl)
        else:
            print(f"DEBUG: Subject grade level missing, keeping existing: {subject.grade_level}", flush=True)
        
        subject.track = request.form.get('track_select')
        
        db.session.commit()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            # Pass a fixed-result lambda to avoid any recursion or shadowing issues
            comp_status = all([getattr(subject, f, None) not in [None, ''] for f in ['name', 'department', 'duration_mins', 'grade_level']])
            return jsonify({
                'success': True, 
                'message': 'Subject updated.', 
                'row_html': render_template('partials/_subject_row.html', s=subject, is_complete=lambda x: comp_status)
            })
            
        flash('Subject updated.')
        return redirect(url_for('admin_subjects'))
    except Exception as e:
        db.session.rollback()
        print(f"ERROR in edit_subject: {str(e)}") # Console debug
        error_msg = f'Error updating subject: {str(e)}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': False, 'error': error_msg}), 500
        flash(error_msg, 'error')
        return redirect(url_for('admin_subjects'))

@app.route('/admin/bulk_edit/<module>', methods=['POST'])
@login_required
def admin_bulk_edit(module):
    if current_user.role != 'admin': return redirect(url_for('index'))
    selected_ids = request.form.getlist('selected_ids')
    if not selected_ids:
        flash("No items selected.")
        return redirect(url_for(f'admin_{module}'))
    
    confirm_overwrite = request.form.get('confirm_overwrite') == 'true'
    
    if module == 'teachers':
        teachers = Teacher.query.filter(Teacher.id.in_(selected_ids)).all()
        max_hours = request.form.get('max_hours_per_day')
        stay_window = request.form.get('stay_window_hours')
        pref_days = request.form.get('preferred_days')
        subjects = request.form.getlist('subjects')
        grade_levels = request.form.getlist('grade_levels')
        is_master = request.form.get('is_master')
        department = request.form.get('department')
        
        for t in teachers:
            if max_hours: t.max_hours_per_day = int(max_hours)
            if stay_window: t.stay_window_hours = int(stay_window)
            if pref_days: t.preferred_days = pref_days
            if subjects: t.subjects = ','.join(subjects)
            if grade_levels: t.grade_levels = ','.join([normalize_gl(g) for g in grade_levels])
            if is_master is not None: t.is_master = (is_master == 'true')
            if department: t.department = department
            
    elif module == 'classrooms':
        rooms = Classroom.query.filter(Classroom.id.in_(selected_ids)).all()
        room_type = request.form.get('room_type')
        building = request.form.get('building')
        for r in rooms:
            if room_type: r.room_type = room_type
            if building: r.building = building
            
    elif module == 'sections':
        sections = Section.query.filter(Section.id.in_(selected_ids)).all()
        adviser_id = request.form.get('adviser_id')
        room_id = request.form.get('room_id')
        track = request.form.get('track_select')
        is_sec_a = request.form.get('is_section_a')
        department = request.form.get('department')
        grade_level = request.form.get('grade_level')
        
        for s in sections:
            if adviser_id: s.adviser_id = int(adviser_id) if adviser_id != 'none' else None
            if room_id: 
                parsed_room_id = int(room_id) if room_id != 'none' else None
                if parsed_room_id:
                    settings_dict = {setting.key: setting.value for setting in Setting.query.all()}
                    inc_shift = get_section_shift(s.department, s.grade_level, settings_dict)
                    is_valid, v_msg = validate_room_assignment(parsed_room_id, s.id, inc_shift, settings_dict)
                    if not is_valid:
                        return jsonify({'success': False, 'msg': f"Validation failed for {s.name}: {v_msg}"}), 400
                s.room_id = parsed_room_id
            if track: s.track = track
            if is_sec_a is not None: s.is_section_a = (is_sec_a == 'true')
            if department: s.department = department
            if grade_level: s.grade_level = normalize_gl(grade_level)
            
    elif module == 'subjects':
        subs = Subject.query.filter(Subject.id.in_(selected_ids)).all()
        req_lab = request.form.get('requires_lab')
        duration = request.form.get('duration_mins')
        meetings = request.form.get('meetings_per_week')
        track = request.form.get('track_select')
        department = request.form.get('department')
        grade_level = request.form.get('grade_level')
        for s in subs:
            if req_lab is not None: s.requires_lab = (req_lab == 'true')
            if duration: s.duration_mins = int(duration)
            if meetings: s.meetings_per_week = int(meetings)
            if grade_level: s.grade_level = normalize_gl(grade_level)
            if track: s.track = track
            if department: s.department = department
            
    db.session.commit()
    flash(f"Updated {len(selected_ids)} records in {module}.")
    return redirect(url_for(f'admin_{module}'))

@app.route('/admin/<module>/import', methods=['POST'])
@login_required
def admin_bulk_import(module):
    if current_user.role != 'admin': return redirect(url_for('index'))
    import csv
    import io
    
    file = request.files.get('file')
    if not file or not file.filename.endswith('.csv'):
        flash('Invalid file. Please upload a CSV file.', 'error')
        return redirect(request.referrer)
        
    stream = io.StringIO(file.stream.read().decode("UTF8"), newline=None)
    csv_reader = csv.DictReader(stream)
    
    added_count = 0
    try:
        for row in csv_reader:
            # Clean keys
            row = {k.strip().lower(): v.strip() for k, v in row.items() if k}
            
            if module == 'teachers':
                name = row.get('name')
                if not name: continue
                teacher = Teacher.query.filter_by(name=name).first() or Teacher(name=name)
                teacher.department = row.get('department', 'JHS')
                teacher.grade_levels = row.get('grade_levels', '')
                teacher.subjects = row.get('subjects', '')
                teacher.max_hours_per_day = int(row.get('max_hours_per_day', 6))
                teacher.stay_window_hours = int(row.get('stay_window_hours', 9))
                teacher.is_master = row.get('is_master', '').lower() in ['yes', 'true', '1']
                if not teacher.id: db.session.add(teacher)
                
            elif module == 'classrooms':
                name = row.get('name')
                if not name: continue
                room = Classroom.query.filter_by(name=name).first() or Classroom(name=name)
                room.room_type = row.get('room_type', 'Room')
                room.building = row.get('building', 'JHS')
                if not room.id: db.session.add(room)
                
            elif module == 'sections':
                name = row.get('name')
                if not name: continue
                section = Section.query.filter_by(name=name).first() or Section(name=name)
                section.department = row.get('department', 'JHS')
                section.grade_level = row.get('grade_level', '7')
                section.track = row.get('track', '')
                
                adv_name = row.get('adviser_name')
                if adv_name:
                    adv = Teacher.query.filter(Teacher.name.ilike(adv_name)).first()
                    if adv: section.adviser_id = adv.id
                
                room_name = row.get('room_name')
                if room_name:
                    rm = Classroom.query.filter(Classroom.name.ilike(room_name)).first()
                    if rm: section.room_id = rm.id
                if not section.id: db.session.add(section)
                
            elif module == 'subjects':
                name = row.get('name')
                if not name: continue
                dept = row.get('department', 'JHS')
                subject = Subject.query.filter_by(name=name, department=dept).first() or Subject(name=name, department=dept)
                subject.duration_mins = int(row.get('duration_mins', 60))
                subject.meetings_per_week = int(row.get('meetings_per_week', 1))
                subject.grade_level = row.get('grade_level', '7')
                subject.track = row.get('track', '')
                subject.requires_lab = row.get('requires_lab', '').lower() in ['yes', 'true', '1']
                if not subject.id: db.session.add(subject)
                
            added_count += 1
            
        db.session.commit()
        flash(f'Successfully imported {added_count} records to {module}.')
    except Exception as e:
        db.session.rollback()
        flash(f'Import error: {str(e)}', 'error')
        
    return redirect(request.referrer)

@app.route('/admin/settings', methods=['GET', 'POST'])
@login_required
def admin_settings():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        # Handle checkboxes and other inputs
        # First clear all previous shift grade levels
        Setting.query.filter(Setting.key.like('jhs_am_grade_%')).delete()
        Setting.query.filter(Setting.key.like('jhs_pm_grade_%')).delete()
        
        # Handle multi-select checkboxes explicitly
        for lst_key in ['active_days', 'jhs_special_days', 'shs_special_days']:
            lst_val = request.form.getlist(lst_key)
            if lst_val:
                lst_str = ','.join(lst_val)
                setting = Setting.query.filter_by(key=lst_key).first()
                if not setting:
                    setting = Setting(key=lst_key, value=lst_str)
                    db.session.add(setting)
                else:
                    setting.value = lst_str

        # If unchecked, they won't be in request.form, so clear them if they were completely unchecked
        for lst_key in ['active_days', 'jhs_special_days', 'shs_special_days', 'jhs_special_enabled', 'shs_special_enabled', 'standard_weekly_load']:
            if not request.form.get(lst_key) and not request.form.getlist(lst_key):
                setting = Setting.query.filter_by(key=lst_key).first()
                if setting:
                    setting.value = ''

        for key, value in request.form.items():
            if key in ['active_days', 'jhs_special_days', 'shs_special_days']: continue # Handled above
            setting = Setting.query.filter_by(key=key).first()
            if not setting:
                setting = Setting(key=key, value=value)
                db.session.add(setting)
            else:
                setting.value = value
        db.session.commit()
        flash('Settings updated.')
        return redirect(url_for('admin_settings'))
    
    settings = {s.key: s.value for s in Setting.query.all()}
    teachers = Teacher.query.all()
    teachers.sort(key=lambda x: natural_sort_key(x.name))
    return render_template('admin_settings.html', settings=settings, teachers=teachers)

@app.route('/admin/change_password', methods=['GET', 'POST'])
@login_required
def admin_change_password():
    if current_user.role != 'admin': return redirect(url_for('index'))
    if request.method == 'POST':
        user_username = request.form.get('user_username')
        new_password = request.form.get('new_password')
        reset_type = request.form.get('reset_type')
        user = User.query.filter_by(username=user_username).first()
        if user:
            if user.is_super_admin and current_user.id != user.id:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'error': 'Unauthorized. Cannot modify other Super Admin credentials.'})
                flash("Unauthorized. Cannot modify other Super Admin credentials.", "error")
                return redirect(url_for('admin_change_password'))

            if reset_type == 'default':
                new_password = get_default_password()
                log_msg = f"Reset password for {user.username} to default"
            else:
                log_msg = f"Changed password for {user.username} manually"
                
            user.password = generate_password_hash(new_password)
            user.password_updated_at = datetime.now()
            db.session.commit()
            log_activity(current_user.username, current_user.role, log_msg, 'Users')
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                formatted_time = user.password_updated_at.strftime('%Y-%m-%d %I:%M %p')
                return jsonify({
                    'success': True,
                    'msg': f"Password updated successfully for {user.username}",
                    'username': user.username,
                    'updated_at': formatted_time
                })
            flash(f"Password changed for user {user.username}")
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': f"User '{user_username}' not found."})
            flash(f"User '{user_username}' not found.")
        return redirect(url_for('admin_change_password'))
    
    users = User.query.filter_by(is_active=True).all()
    users.sort(key=lambda x: natural_sort_key(x.username))
    sections = Section.query.all()
    sections.sort(key=lambda x: natural_sort_key(x.name))
    return render_template('admin_password.html', users=users, sections=sections, default_password=get_default_password())

@app.route('/admin/default_password', methods=['POST'])
@login_required
def admin_default_password():
    if current_user.role != 'admin': return jsonify({'success': False, 'error': 'Unauthorized'}), 403
    
    action = request.form.get('action')
    if action == 'update_default':
        new_default = request.form.get('new_default')
        if not new_default: return jsonify({'success': False, 'error': 'Password cannot be empty'})
        setting = Setting.query.filter_by(key='default_password').first()
        if not setting:
            setting = Setting(key='default_password', value=new_default)
            db.session.add(setting)
        else:
            setting.value = new_default
        db.session.commit()
        log_activity(current_user.username, current_user.role, "Updated Default Password Setting", "Settings")
        return jsonify({'success': True, 'msg': 'Default password updated successfully.'})
        
    elif action == 'bulk_apply':
        apply_target = request.form.get('apply_target')
        default_pw = get_default_password()
        
        query = User.query
        if apply_target == 'teachers':
            query = query.filter_by(role='teacher')
            target_name = "Teachers"
        elif apply_target == 'students':
            query = query.filter_by(role='student')
            target_name = "Students"
        elif apply_target == 'admins':
            query = query.filter_by(role='admin')
            target_name = "Admins"
        else:
            return jsonify({'success': False, 'error': 'Invalid target'})
            
        users = query.all()
        count = 0
        for u in users:
            if getattr(u, 'is_super_admin', False): continue
            u.password = generate_password_hash(default_pw)
            u.password_updated_at = datetime.now()
            count += 1
            
        db.session.commit()
        log_activity(current_user.username, current_user.role, f"Bulk Reset Passwords to Default ({target_name})", "Users")
        return jsonify({'success': True, 'msg': f'Successfully applied default password to {count} {target_name}.'})
        
    return jsonify({'success': False, 'error': 'Invalid action'})

@app.route('/teacher/change_password', methods=['GET', 'POST'])
@login_required
def teacher_change_password():
    if current_user.role != 'teacher': 
        if request.method == 'POST':
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        current_pw = request.form.get('current_password')
        new_pw = request.form.get('new_password')
        
        if not check_password_hash(current_user.password, current_pw):
            return jsonify({'success': False, 'error': 'Incorrect current password.'})
            
        import re
        if len(new_pw) < 8 or not re.search('[a-zA-Z]', new_pw) or not re.search('[0-9]', new_pw):
            return jsonify({'success': False, 'error': 'New password must be at least 8 characters and contain both letters and numbers.'})
            
        current_user.password = generate_password_hash(new_pw)
        current_user.password_updated_at = datetime.now()
        db.session.commit()
        
        log_activity(current_user.username, current_user.role, "Changed Own Password", "Authentication")
        return jsonify({'success': True, 'msg': 'Password updated successfully!'})
        
    return render_template('teacher_password.html')

@app.route('/admin/create_user', methods=['POST'])
@login_required
def admin_create_user():
    if current_user.role != 'admin' or not current_user.is_super_admin:
        return jsonify({'success': False, 'error': 'Unauthorized. Only Super Admin can create user accounts.'}), 403
        
    username = request.form.get('username', '').strip()
    role = request.form.get('role')
    password_mode = request.form.get('password_mode', 'auto')
    password = request.form.get('password', '').strip()
    section_id = request.form.get('section_id')
    teacher_dept = request.form.get('teacher_dept', 'JHS')
    
    if not username:
        return jsonify({'success': False, 'error': 'Username is required.'}), 400
        
    if role not in ['admin', 'teacher', 'student']:
        return jsonify({'success': False, 'error': 'Invalid role choice.'}), 400
        
    # Check duplicate username
    existing = User.query.filter_by(username=username).first()
    if existing:
        return jsonify({'success': False, 'error': f"Username '{username}' is already taken."}), 400
        
    if password_mode == 'auto':
        password = get_default_password()
    else:
        if not password:
            return jsonify({'success': False, 'error': 'Password is required for manual mode.'}), 400
        import re
        if len(password) < 8 or not re.search('[a-zA-Z]', password) or not re.search('[0-9]', password):
            return jsonify({'success': False, 'error': 'Password must be at least 8 characters and contain both letters and numbers.'}), 400
            
    try:
        related_id = None
        
        if role == 'teacher':
            # Create corresponding Teacher record
            teacher = Teacher(
                name=username,
                department=teacher_dept,
                grade_levels='7',
                max_hours_per_day=6,
                stay_window_hours=9,
                is_master=False,
                handle_sec_a=False,
                preferred_days='Mon-Fri',
                subjects=''
            )
            db.session.add(teacher)
            db.session.flush() # get teacher.id
            related_id = teacher.id
            
        elif role == 'student':
            if section_id:
                related_id = int(section_id)
            else:
                first_section = Section.query.first()
                if first_section:
                    related_id = first_section.id
                else:
                    # Create placeholder section
                    sec = Section(name=f"{username}_Sec", department="JHS", grade_level="7", adviser_id=None, room_id=None)
                    db.session.add(sec)
                    db.session.flush()
                    related_id = sec.id
                    
        new_user = User(
            username=username,
            password=generate_password_hash(password),
            role=role,
            related_id=related_id,
            is_active=True
        )
        db.session.add(new_user)
        db.session.commit()
        
        # Log this user creation activity
        log_activity(
            actor_username=current_user.username,
            role=current_user.role,
            action=f"Created User Account: {username} ({role.capitalize()})",
            module='Users'
        )
        
        return jsonify({
            'success': True,
            'msg': f"User account successfully created for {username}!",
            'username': username,
            'role': role,
            'user_id': new_user.id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': f"Failed to create user: {str(e)}"}), 500

@app.route('/admin/users/edit/<int:user_id>', methods=['POST'])
@login_required
def admin_edit_user(user_id):
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Unauthorized access. Only administrators can edit usernames.'}), 403
        
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'error': 'User account not found.'}), 404
        
    # Restriction: Regular Admin cannot edit Super Admin accounts
    if getattr(user, 'is_super_admin', False) and not current_user.is_super_admin:
        return jsonify({'success': False, 'error': 'Unauthorized. Regular administrators cannot edit Super Admin accounts.'}), 403
        
    new_username = request.form.get('username', '').strip()
    if not new_username:
        return jsonify({'success': False, 'error': 'New username cannot be empty.'}), 400
        
    old_username = user.username
    if new_username == old_username:
        return jsonify({'success': False, 'error': 'New username must be different.'}), 400
        
    # Check duplicate username
    existing = User.query.filter(User.username == new_username, User.id != user_id).first()
    if existing:
        return jsonify({'success': False, 'error': f"Username '{new_username}' is already taken."}), 400
        
    try:
        user.username = new_username
        
        # If the user is linked to a teacher, update the Teacher's name
        if user.role == 'teacher' and user.related_id:
            teacher = Teacher.query.get(user.related_id)
            if teacher:
                teacher.name = new_username
                
        db.session.commit()
        
        # Log this administrative edit activity
        log_activity(
            actor_username=current_user.username,
            role=current_user.role,
            action=f"Edited User Account Username (Old: {old_username} -> New: {new_username})",
            module='Users'
        )
        
        return jsonify({
            'success': True,
            'msg': f"Username successfully updated from '{old_username}' to '{new_username}'!",
            'username': new_username,
            'user_id': user.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': f"Failed to update username: {str(e)}"}), 500

@app.route('/admin/users/delete/<int:user_id>', methods=['POST'])
@login_required
def admin_delete_user(user_id):
    if current_user.role != 'admin':
        return jsonify({'success': False, 'error': 'Unauthorized access. Only administrators can delete accounts.'}), 403
        
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'error': 'User account not found.'}), 404
        
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'You cannot delete your own account.'}), 400
        
    # Restriction: Regular Admin cannot delete Super Admin accounts
    if getattr(user, 'is_super_admin', False) and not current_user.is_super_admin:
        return jsonify({'success': False, 'error': 'Unauthorized. Regular administrators cannot delete Super Admin accounts.'}), 403
        
    try:
        target_username = user.username
        
        # Soft delete: de-activate account
        user.is_active = False
        db.session.commit()
        
        # Log this administrative delete activity
        log_activity(
            actor_username=current_user.username,
            role=current_user.role,
            action=f"Soft Deleted User Account: {target_username}",
            module='Users'
        )
        
        return jsonify({
            'success': True,
            'msg': f"User account '{target_username}' successfully deleted (deactivated).",
            'username': target_username,
            'user_id': user.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': f"Failed to delete user account: {str(e)}"}), 500

@app.route('/admin/activity_logs', methods=['GET'])
@login_required
def admin_activity_logs():
    if current_user.role != 'admin':
        return redirect(url_for('index'))
        
    # Get filters
    date_filter = request.args.get('date_filter', '30') # default 30 days
    action_type = request.args.get('action_type', 'all')
    user_filter = request.args.get('user_filter', 'all')
    search_query = request.args.get('search', '').strip()
    
    # Base Query
    query = ActivityLog.query
    
    # Access control: Admin sees only their own; Super Admin sees all!
    if not current_user.is_super_admin:
        query = query.filter(ActivityLog.actor_username == current_user.username)
    else:
        if user_filter != 'all' and user_filter != '' and user_filter is not None:
            query = query.filter(ActivityLog.actor_username == user_filter)
            
    # Date filtering
    from datetime import timedelta
    now = datetime.now()
    if date_filter == '7':
        query = query.filter(ActivityLog.timestamp >= now - timedelta(days=7))
    elif date_filter == '30':
        query = query.filter(ActivityLog.timestamp >= now - timedelta(days=30))
    elif date_filter == '90':
        query = query.filter(ActivityLog.timestamp >= now - timedelta(days=90))
    
    # Action type filter
    if action_type != 'all' and action_type != '' and action_type is not None:
        query = query.filter(ActivityLog.action.ilike(f"%{action_type}%"))
        
    # Search query
    if search_query:
        query = query.filter(
            (ActivityLog.actor_username.ilike(f"%{search_query}%")) |
            (ActivityLog.action.ilike(f"%{search_query}%")) |
            (ActivityLog.module.ilike(f"%{search_query}%"))
        )
        
    logs = query.order_by(ActivityLog.timestamp.desc()).all()
    
    # Get all unique actors for the filter dropdown (Super Admin only)
    unique_users = []
    if current_user.is_super_admin:
        unique_users = [u[0] for u in db.session.query(ActivityLog.actor_username).distinct().all() if u[0]]
        unique_users.sort()
        
    return render_template('admin_activity_logs.html', 
                           logs=logs, 
                           unique_users=unique_users,
                           current_date_filter=date_filter,
                           current_action_type=action_type,
                           current_user_filter=user_filter,
                           search_query=search_query)


def get_global_schedule_bounds():
    """Returns the normalized absolute min_start (HH:MM) and max_end (HH:MM) across all shifts."""
    start_keys = ['jhs_am_start', 'jhs_pm_start', 'shs_start']
    end_keys = ['jhs_am_end', 'jhs_pm_end', 'shs_end']
    
    all_keys = start_keys + end_keys
    settings = {s.key: s.value for s in Setting.query.filter(Setting.key.in_(all_keys)).all()}
    
    # Defaults
    default_vals = {
        'jhs_am_start': '06:00', 'jhs_am_end': '12:45',
        'jhs_pm_start': '13:00', 'jhs_pm_end': '19:45',
        'shs_start': '07:00', 'shs_end': '17:00'
    }
    
    starts_m = []
    for k in start_keys:
        val = settings.get(k, default_vals.get(k))
        try:
            starts_m.append(time_to_min(val))
        except: continue
        
    ends_m = []
    for k in end_keys:
        val = settings.get(k, default_vals.get(k))
        try:
            ends_m.append(time_to_min(val))
        except: continue
            
    if not starts_m: starts_m = [time_to_min('06:00')]
    if not ends_m: ends_m = [time_to_min('19:45')]
        
    abs_min = min(starts_m)
    abs_max = max(ends_m)
    
    # Normalize to 5-minute grid (Floor start, Ceil end)
    norm_min = (abs_min // 5) * 5
    norm_max = ((abs_max + 4) // 5) * 5
    
    # Ensure at least one 5-min slot
    if norm_max <= norm_min:
        norm_max = norm_min + 5
        
    return min_to_time(norm_min), min_to_time(norm_max)

def get_section_shift(dept, grade_level, settings_dict):
    if dept == 'SHS':
        return 'FULL_DAY'
    
    gl = str(grade_level).upper().replace('GRADE', '').strip()
    is_am = settings_dict.get(f'jhs_am_grade_{gl}') in ['active', 'on']
    is_pm = settings_dict.get(f'jhs_pm_grade_{gl}') in ['active', 'on']
    
    if is_am and is_pm: return 'FULL_DAY'
    elif is_pm and not is_am: return 'PM'
    else: return 'AM'

def validate_room_assignment(room_id, incoming_section_id, incoming_shift, settings_dict):
    if not room_id:
        return True, ""
        
    existing_sections = Section.query.filter(Section.room_id == room_id).all()
    
    for existing in existing_sections:
        if incoming_section_id and existing.id == incoming_section_id:
            continue
            
        existing_shift = get_section_shift(existing.department, existing.grade_level, settings_dict)
        
        if incoming_shift == 'FULL_DAY' or existing_shift == 'FULL_DAY':
            return False, f"Room is blocked by a FULL_DAY section ({existing.name})."
            
        if incoming_shift == existing_shift:
            return False, f"Room is already assigned to a {existing_shift} section ({existing.name})."
            
    return True, ""

def get_shift_bounds(dept, grade_level, settings_dict):
    if dept == 'SHS':
        return settings_dict.get('shs_start', '07:00'), settings_dict.get('shs_end', '17:00')
    else:
        gl = str(grade_level).upper().replace('GRADE', '').strip()
        # Recognize both 'active' (API) and 'on' (HTML Form) as truthy
        is_am = settings_dict.get(f'jhs_am_grade_{gl}') in ['active', 'on']
        is_pm = settings_dict.get(f'jhs_pm_grade_{gl}') in ['active', 'on']
        if is_pm and not is_am:
            return settings_dict.get('jhs_pm_start', '12:00'), settings_dict.get('jhs_pm_end', '18:00')
        else:
            return settings_dict.get('jhs_am_start', '06:00'), settings_dict.get('jhs_am_end', '12:00')

def prepare_schedule_grid(entity_id, view_type, schedules, force_start=None, force_end=None):
    # Fetch active days from settings
    active_days_setting = Setting.query.filter_by(key='active_days').first()
    if active_days_setting:
        days = active_days_setting.value.split(',')
    else:
        days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
        
    settings_dict = {s.key: s.value for s in Setting.query.all()}

    abs_min_m = time_to_min('18:00')
    abs_max_m = time_to_min('06:00')
    
    entity_shift_s = '06:00'
    entity_shift_e = '18:00'

    if view_type == 'teacher':
        entity_schedules = [s for s in schedules if int(s.teacher_id) == int(entity_id)]
        teacher = Teacher.query.get(entity_id)
        break_times = []
        if teacher:
            assigned_sec_ids = {s.section_id for s in entity_schedules}
            if not assigned_sec_ids:
                abs_min_m = time_to_min('06:00')
                abs_max_m = time_to_min('18:00')
            else:
                for sid in assigned_sec_ids:
                    sec = Section.query.get(sid)
                    if sec:
                        try:
                            st, en = get_shift_bounds(sec.department, sec.grade_level, settings_dict)
                            st_m, en_m = time_to_min(st), time_to_min(en)
                            if st_m < abs_min_m: abs_min_m = st_m
                            if en_m > abs_max_m: abs_max_m = en_m
                        except:
                            continue
                
                # Update shift bounds for masking
                entity_shift_s = min_to_time(abs_min_m)
                entity_shift_e = min_to_time(abs_max_m)
                
                # FINAL EXPANSION: Ensure actual schedule times are also covered
                for sch in entity_schedules:
                    sch_s = time_to_min(sch.start_time)
                    sch_e = time_to_min(sch.end_time)
                    if sch_s < abs_min_m: abs_min_m = sch_s
                    if sch_e > abs_max_m: abs_max_m = sch_e
            
            if 'JHS' in teacher.department:
                if settings_dict.get('jhs_am_break_start') and settings_dict.get('jhs_am_break_end'):
                    break_times.append((settings_dict['jhs_am_break_start'], settings_dict['jhs_am_break_end']))
                if settings_dict.get('jhs_pm_break_start') and settings_dict.get('jhs_pm_break_end'):
                    break_times.append((settings_dict['jhs_pm_break_start'], settings_dict['jhs_pm_break_end']))
            if 'SHS' in teacher.department:
                b_s = settings_dict.get('shs_break_start', '09:30')
                b_e = settings_dict.get('shs_break_end', '10:00')
                l_s = settings_dict.get('shs_lunch_start', '12:00')
                l_e = settings_dict.get('shs_lunch_end', '13:00')
                break_times.append((b_s, b_e))
                break_times.append((l_s, l_e))
    elif view_type == 'section':
        entity_schedules = [s for s in schedules if s.section_id == entity_id]
        section = Section.query.get(entity_id)
        break_times = []
        if section:
            st, en = get_shift_bounds(section.department, section.grade_level, settings_dict)
            abs_min_m = time_to_min(st)
            abs_max_m = time_to_min(en)
            entity_shift_s = st
            entity_shift_e = en
            
            if section.department == 'SHS':
                b_s = settings_dict.get('shs_break_start', '09:30')
                b_e = settings_dict.get('shs_break_end', '10:00')
                l_s = settings_dict.get('shs_lunch_start', '12:00')
                l_e = settings_dict.get('shs_lunch_end', '13:00')
                break_times.append((b_s, b_e))
                break_times.append((l_s, l_e))
            else:
                gl = str(section.grade_level).upper().replace('GRADE', '').strip()
                is_am = settings_dict.get(f"jhs_am_grade_{gl}") in ['active', 'on']
                is_pm = settings_dict.get(f"jhs_pm_grade_{gl}") in ['active', 'on']
                shift = 'PM' if (is_pm and not is_am) else 'AM'
                prefix = 'jhs_am' if shift == 'AM' else 'jhs_pm'
                default_s = '09:00' if shift == 'AM' else '15:00'
                default_e = '09:30' if shift == 'AM' else '15:30'
                break_times.append((settings_dict.get(f'{prefix}_break_start', default_s), settings_dict.get(f'{prefix}_break_end', default_e)))
    elif view_type == 'classroom':
        entity_schedules = [s for s in schedules if s.room_id == entity_id]
        classroom = Classroom.query.get(entity_id)
        break_times = []
        
        # Determine room shift range from assignments
        assigned_sec_ids = {s.section_id for s in entity_schedules}
        if assigned_sec_ids:
            rm_min_m = time_to_min('18:00')
            rm_max_m = time_to_min('06:00')
            assigned_depts = set()
            for sid in assigned_sec_ids:
                sec = Section.query.get(sid)
                if sec:
                    assigned_depts.add(sec.department)
                    st, en = get_shift_bounds(sec.department, sec.grade_level, settings_dict)
                    st_m, en_m = time_to_min(st), time_to_min(en)
                    if st_m < rm_min_m: rm_min_m = st_m
                    if en_m > rm_max_m: rm_max_m = en_m
            abs_min_m, abs_max_m = rm_min_m, rm_max_m
        else:
            st, en = ('06:00', '12:00')
            if classroom and classroom.building == 'SHS':
                st, en = ('07:00', '17:00')
            abs_min_m, abs_max_m = time_to_min(st), time_to_min(en)
        
        entity_shift_s = min_to_time(abs_min_m)
        entity_shift_e = min_to_time(abs_max_m)

        # Break injection for Classroom view:
        # - `break_times` is used on non-special days (special-day logic handled below per-day).
        # - Previously, SHS breaks were included but JHS breaks were not, causing missing breaks for JHS classrooms.
        if assigned_sec_ids:
            # Prefer actual assigned sections' departments when available.
            has_jhs = 'JHS' in assigned_depts
            has_shs = 'SHS' in assigned_depts
        else:
            # Fall back to classroom building when no schedules exist.
            has_jhs = bool(classroom and classroom.building in ['JHS', 'Both'])
            has_shs = bool(classroom and classroom.building in ['SHS', 'Both'])
        
        if has_jhs:
            if settings_dict.get('jhs_am_break_start') and settings_dict.get('jhs_am_break_end'):
                break_times.append((settings_dict['jhs_am_break_start'], settings_dict['jhs_am_break_end']))
            if settings_dict.get('jhs_pm_break_start') and settings_dict.get('jhs_pm_break_end'):
                break_times.append((settings_dict['jhs_pm_break_start'], settings_dict['jhs_pm_break_end']))
        
        if has_shs:
            b_s = settings_dict.get('shs_break_start', '09:30')
            b_e = settings_dict.get('shs_break_end', '10:00')
            l_s = settings_dict.get('shs_lunch_start', '12:00')
            l_e = settings_dict.get('shs_lunch_end', '13:00')
            break_times.append((b_s, b_e))
            break_times.append((l_s, l_e))
    else:
        entity_schedules = schedules
        break_times = []
        abs_min_m = time_to_min('06:00')
        abs_max_m = time_to_min('18:00')
        entity_shift_s, entity_shift_e = '06:00', '18:00'

    if force_start is not None:
        abs_min_m = time_to_min(force_start)
    if force_end is not None:
        abs_max_m = time_to_min(force_end)

    # Ensure valid bounds (fallback to 6am-6pm if logic fails or bounds are equal)
    if abs_min_m >= abs_max_m:
        abs_min_m = time_to_min('06:00')
        abs_max_m = time_to_min('18:00')

    # 5. Build Final Time Slots
    # Consistently floor start and ceil end to 5-minute grid
    min_m = (abs_min_m // 5) * 5
    max_m = ((abs_max_m + 4) // 5) * 5
    
    if max_m <= min_m:
        max_m = min_m + 5
    
    time_slots = [min_to_time(m) for m in range(min_m, max_m, 5)]
    
    grid = {slot: {day: None for day in days} for slot in time_slots}
    spans = {slot: {day: 1 for day in days} for slot in time_slots}
    occupied = {slot: {day: False for day in days} for slot in time_slots}
    break_ends = {slot: {day: None for day in days} for slot in time_slots}

    # Map break times to slots per day
    day_breaks = {d: set() for d in days}
    for d in days:
        current_breaks = []
        if view_type == 'section' and section:
            dept = section.department
            spec_enabled = settings_dict.get(f'{dept.lower()}_special_enabled') == 'yes'
            spec_days = settings_dict.get(f'{dept.lower()}_special_days', '').split(',')
            is_special = spec_enabled and d in spec_days
            
            if is_special:
                if dept == 'JHS':
                    gl_clean = str(section.grade_level).upper().replace('GRADE', '').strip()
                    is_am_sec = settings_dict.get(f"jhs_am_grade_{gl_clean}") in ['active', 'on']
                    is_pm_sec = settings_dict.get(f"jhs_pm_grade_{gl_clean}") in ['active', 'on']
                    sec_shift = 'PM' if (is_pm_sec and not is_am_sec) else 'AM'
                    
                    prefix = 'jhs_am' if sec_shift == 'AM' else 'jhs_pm'
                    def_s = '08:40' if sec_shift == 'AM' else '14:40'
                    def_e = '09:10' if sec_shift == 'AM' else '15:10'
                    s_start = settings_dict.get(f'{prefix}_special_break_start', def_s)
                    s_end = settings_dict.get(f'{prefix}_special_break_end', def_e)
                    if s_start and s_end: current_breaks.append((s_start, s_end))
                else:
                    # SHS special breaks
                    s_s = settings_dict.get('shs_special_break_start', '09:30')
                    s_e = settings_dict.get('shs_special_break_end', '10:00')
                    current_breaks.append((s_s, s_e))
                    
                    # SHS special days should still have lunch if whole day
                    l_s = settings_dict.get('shs_lunch_start', '12:00')
                    l_e = settings_dict.get('shs_lunch_end', '13:00')
                    current_breaks.append((l_s, l_e))
            else:
                current_breaks = break_times
        else:
            # For teacher/room/admin views, we check if special mode is active for EITHER department on this day
            jhs_spec = settings_dict.get('jhs_special_enabled') == 'yes' and d in settings_dict.get('jhs_special_days', '').split(',')
            shs_spec = settings_dict.get('shs_special_enabled') == 'yes' and d in settings_dict.get('shs_special_days', '').split(',')
            
            if jhs_spec or shs_spec:
                # Add JHS special breaks if applicable
                if jhs_spec:
                    if settings_dict.get('jhs_am_special_break_start') and settings_dict.get('jhs_am_special_break_end'):
                        current_breaks.append((settings_dict['jhs_am_special_break_start'], settings_dict['jhs_am_special_break_end']))
                    if settings_dict.get('jhs_pm_special_break_start') and settings_dict.get('jhs_pm_special_break_end'):
                        current_breaks.append((settings_dict['jhs_pm_special_break_start'], settings_dict['jhs_pm_special_break_end']))
                # Add SHS special breaks if applicable
                if shs_spec:
                    if settings_dict.get('shs_special_break_start') and settings_dict.get('shs_special_break_end'):
                        current_breaks.append((settings_dict['shs_special_break_start'], settings_dict['shs_special_break_end']))
                
                # If NOT in special mode for a department, we should still include its regular breaks
                if not jhs_spec:
                    if settings_dict.get('jhs_am_break_start') and settings_dict.get('jhs_am_break_end'):
                        current_breaks.append((settings_dict['jhs_am_break_start'], settings_dict['jhs_am_break_end']))
                    if settings_dict.get('jhs_pm_break_start') and settings_dict.get('jhs_pm_break_end'):
                        current_breaks.append((settings_dict['jhs_pm_break_start'], settings_dict['jhs_pm_break_end']))
                if not shs_spec:
                    s_b_s = settings_dict.get('shs_break_start', '09:30')
                    s_b_e = settings_dict.get('shs_break_end', '10:00')
                    l_s = settings_dict.get('shs_lunch_start', '12:00')
                    l_e = settings_dict.get('shs_lunch_end', '13:00')
                    current_breaks.append((s_b_s, s_b_e))
                    current_breaks.append((l_s, l_e))
            else:
                current_breaks = break_times
            
        for b_start, b_end in current_breaks:
            bs_m = time_to_min(b_start)
            be_m = time_to_min(b_end)
            for m in range(bs_m, be_m, 5):
                day_breaks[d].add(min_to_time(m))

    for sch in entity_schedules:
        start_m = time_to_min(sch.start_time)
        end_m = time_to_min(sch.end_time)
        num_slots = (end_m - start_m) // 5
        sch_slots = [min_to_time(start_m + (i * 5)) for i in range(num_slots)]
        
        if sch_slots:
            first_slot = sch_slots[0]
            if first_slot in grid:
                if sch.day_of_week in grid[first_slot]:
                    grid[first_slot][sch.day_of_week] = sch
                    spans[first_slot][sch.day_of_week] = len(sch_slots)
            for slot in sch_slots[1:]:
                if slot in occupied:
                    occupied[slot][sch.day_of_week] = True
    
    # Mark breaks in the grid where NOT occupied by a class
    for day in days:
        i = 0
        while i < len(time_slots):
            slot = time_slots[i]
            if slot in day_breaks[day] and grid[slot][day] is None and not occupied[slot][day]:
                # Find how many consecutive break slots we have
                start_idx = i
                while i < len(time_slots) and time_slots[i] in day_breaks[day] and grid[time_slots[i]][day] is None and not occupied[time_slots[i]][day]:
                    i += 1
                
                num_slots = i - start_idx
                if num_slots > 0:
                    first_slot = time_slots[start_idx]
                    l_s = settings_dict.get('shs_lunch_start', '12:00')
                    is_lunch = False
                    # Check if any slot in this block matches lunch start
                    for j in range(start_idx, i):
                        if time_slots[j] == l_s:
                            is_lunch = True
                            break
                    
                    grid[first_slot][day] = 'LUNCH' if is_lunch else 'BREAK'
                    spans[first_slot][day] = num_slots
                    for j in range(start_idx + 1, i):
                        occupied[time_slots[j]][day] = True
            else:
                i += 1
                    
    slot_ranges = {slot: min_to_time(time_to_min(slot) + 5) for slot in time_slots}
                
    return {
        'grid': grid,
        'spans': spans,
        'occupied': occupied,
        'days': days,
        'time_slots': time_slots,
        'slot_ranges': slot_ranges,
        'break_ends': break_ends,
        'shift_start_m': time_to_min(entity_shift_s),
        'shift_end_m': time_to_min(entity_shift_e)
    }

def prepare_condensed_grid(entity_id, view_type, schedules, force_start=None, force_end=None):
    # Fetch active days from settings
    active_days_setting = Setting.query.filter_by(key='active_days').first()
    if active_days_setting:
        days = active_days_setting.value.split(',')
    else:
        days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    
    settings_dict = {s.key: s.value for s in Setting.query.all()}
    
    # 1. Collect all schedules for this entity
    if view_type == 'teacher':
        entity_schedules = [s for s in schedules if s.teacher_id == entity_id]
        entity = Teacher.query.get(entity_id)
        department = entity.department if entity else 'Both'
    elif view_type == 'section':
        entity_schedules = [s for s in schedules if s.section_id == entity_id]
        entity = Section.query.get(entity_id)
        department = entity.department if entity else 'JHS'
    elif view_type == 'classroom':
        entity_schedules = [s for s in schedules if s.room_id == entity_id]
        c = Classroom.query.get(entity_id)
        department = c.building if c else 'Both'
    else:
        entity_schedules = schedules
        department = 'Both'

    # 2. Collect all unique time bounds (start and end times)
    time_bounds = set()
    # 3. Determine natural shift bounds for masking
    if view_type == 'section' and entity:
        st_s, st_e = get_shift_bounds(entity.department, entity.grade_level, settings_dict)
        shift_s_m, shift_e_m = time_to_min(st_s), time_to_min(st_e)
    else:
        # Union-based shift detection
        if entity_schedules:
            shift_s_m = min(time_to_min(s.start_time) for s in entity_schedules)
            shift_e_m = max(time_to_min(s.end_time) for s in entity_schedules)
        else:
            shift_s_m, shift_e_m = time_to_min('07:00'), time_to_min('17:00')

    for sch in entity_schedules:
        time_bounds.add(sch.start_time)
        time_bounds.add(sch.end_time)
    
    # 3. Add break times as bounds (Day-Aware)
    day_breaks = {d: [] for d in days}
    
    for d in days:
        # Determine breaks for this day based on department and mode
        current_day_breaks = []
        if 'JHS' in department or department == 'Both':
            spec_enabled = settings_dict.get('jhs_special_enabled') == 'yes'
            spec_days = settings_dict.get('jhs_special_days', '').split(',')
            is_special = spec_enabled and d in spec_days
            
            if is_special:
                # Need to handle both shifts if department is 'Both' or in non-section view
                s_am_start = settings_dict.get('jhs_am_special_break_start')
                s_am_end = settings_dict.get('jhs_am_special_break_end')
                s_pm_start = settings_dict.get('jhs_pm_special_break_start')
                s_pm_end = settings_dict.get('jhs_pm_special_break_end')
                
                if s_am_start and s_am_end: current_day_breaks.append((s_am_start, s_am_end))
                if s_pm_start and s_pm_end: current_day_breaks.append((s_pm_start, s_pm_end))
            else:
                # Regular JHS Breaks
                j_am_s = settings_dict.get('jhs_am_break_start', '09:00')
                j_am_e = settings_dict.get('jhs_am_break_end', '09:30')
                j_pm_s = settings_dict.get('jhs_pm_break_start', '15:00')
                j_pm_e = settings_dict.get('jhs_pm_break_end', '15:30')
                current_day_breaks.append((j_am_s, j_am_e))
                current_day_breaks.append((j_pm_s, j_pm_e))
                    
        if 'SHS' in department or department == 'Both':
            spec_enabled = settings_dict.get('shs_special_enabled') == 'yes'
            spec_days = settings_dict.get('shs_special_days', '').split(',')
            is_special = spec_enabled and d in spec_days
            
            if is_special:
                s_start = settings_dict.get('shs_special_break_start', '09:30')
                s_end = settings_dict.get('shs_special_break_end', '10:00')
                current_day_breaks.append((s_start, s_end))
                # Preserving lunch on special days
                l_s = settings_dict.get('shs_lunch_start', '12:00')
                l_e = settings_dict.get('shs_lunch_end', '13:00')
                current_day_breaks.append((l_s, l_e))
            else:
                s_b_s = settings_dict.get('shs_break_start', '09:30')
                s_b_e = settings_dict.get('shs_break_end', '10:00')
                current_day_breaks.append((s_b_s, s_b_e))
                l_s = settings_dict.get('shs_lunch_start', '12:00')
                l_e = settings_dict.get('shs_lunch_end', '13:00')
                current_day_breaks.append((l_s, l_e))
        
        day_breaks[d] = current_day_breaks
        for b_start, b_end in current_day_breaks:
            time_bounds.add(b_start)
            time_bounds.add(b_end)
    
    # Ensure min/max bounds are present if no classes scheduled
    if not time_bounds:
        time_bounds.add('07:00'); time_bounds.add('17:00')
        
    # Ensure shift bounds are in time_bounds to create perfect boundaries
    time_bounds.add(min_to_time(shift_s_m))
    time_bounds.add(min_to_time(shift_e_m))
    
    if force_start: time_bounds.add(force_start)
    if force_end: time_bounds.add(force_end)

    sorted_bounds = sorted(list(time_bounds), key=time_to_min)
    
    # 4. Create condensed time slots
    time_slots = []
    slot_ranges = {}
    for i in range(len(sorted_bounds) - 1):
        start = sorted_bounds[i]
        end = sorted_bounds[i+1]
        s_m = time_to_min(start)
        e_m = time_to_min(end)
        
        # Only include slots that fall completely within the shift bounds
        if s_m >= shift_s_m and e_m <= shift_e_m and s_m < e_m:
            time_slots.append(start)
            slot_ranges[start] = end
    
    # 5. Initialize grid
    grid = {slot: {day: None for day in days} for slot in time_slots}
    spans = {slot: {day: 1 for day in days} for slot in time_slots}
    occupied = {slot: {day: False for day in days} for slot in time_slots}
    
    # 6. Map schedules to condensed slots
    for sch in entity_schedules:
        s_m = time_to_min(sch.start_time)
        e_m = time_to_min(sch.end_time)
        # Use < e_m to exclude the slot starting exactly AT the boundary time
        covered_slots = [slot for slot in time_slots if time_to_min(slot) >= s_m and time_to_min(slot) < e_m]
        
        if covered_slots:
            first_slot = covered_slots[0]
            grid[first_slot][sch.day_of_week] = sch
            spans[first_slot][sch.day_of_week] = len(covered_slots)
            for slot in covered_slots[1:]:
                occupied[slot][sch.day_of_week] = True
    
    # 7. Map breaks (Day-Aware Mapping)
    for day in days:
        for b_start, b_end in day_breaks[day]:
            bs_m = time_to_min(b_start)
            be_m = time_to_min(b_end)
            covered_slots = [slot for slot in time_slots if time_to_min(slot) >= bs_m and time_to_min(slot) < be_m]
            
            if covered_slots:
                first_slot = covered_slots[0]
                if grid[first_slot][day] is None and not occupied[first_slot][day]:
                    all_free = True
                    for slot in covered_slots:
                        if grid[slot][day] is not None or occupied[slot][day]:
                            all_free = False; break
                    if all_free:
                        l_s = settings_dict.get('shs_lunch_start', '12:00')
                        is_lunch = (b_start == l_s)
                        grid[first_slot][day] = 'LUNCH' if is_lunch else 'BREAK'
                        spans[first_slot][day] = len(covered_slots)
                        for slot in covered_slots[1:]:
                            occupied[slot][day] = True
                    
    is_shared_am_pm = False
    if view_type == 'classroom':
        am_sch = [s for s in entity_schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'AM']
        pm_sch = [s for s in entity_schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'PM']
        if am_sch and pm_sch:
            is_shared_am_pm = True

    return {
        'grid': grid,
        'spans': spans,
        'occupied': occupied,
        'days': days,
        'time_slots': time_slots,
        'slot_ranges': slot_ranges,
        'shift_start_m': shift_s_m,
        'shift_end_m': shift_e_m,
        'is_shared_am_pm': is_shared_am_pm
    }

@app.route('/admin/schedule/generate', methods=['POST'])
@login_required
def admin_generate_schedule():
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.get_json() or {}
    phase = data.get('phase', 'all')
    
    global scheduler_status
    if scheduler_status['running']:
        return jsonify({'error': 'Generation already in progress.'}), 400
        
    phase_map = {
        'jhs': 'JHS Schedule',
        'shs': 'SHS Schedule',
        'all': 'Full Regeneration'
    }
    action_str = f"Initiated {phase_map.get(phase, phase.upper())}"
    
    log_activity(
        actor_username=current_user.username,
        role=current_user.role,
        action=action_str,
        module='Schedule'
    )

    def run_wrapper(app_context, phase):
        with app_context:
            try:
                global scheduler_status
                scheduler_status['running'] = True
                scheduler_status['progress'] = 0
                scheduler_status['phase'] = phase
                scheduler_status['message'] = f'Starting {phase.upper()} AI engine...'
                scheduler_status['cancel_requested'] = False
                
                from scheduler import generate_schedule
                
                def progress_cb(p, m):
                    scheduler_status['progress'] = p
                    if ' | ' in m:
                        parts = m.split(' | ')
                        scheduler_status['message'] = parts[0]
                        for part in parts[1:]:
                            if ':' in part:
                                key, val = part.split(':')
                                scheduler_status[key] = val
                    else:
                        scheduler_status['message'] = m
                
                def stop_check():
                    return scheduler_status.get('cancel_requested', False)
                
                success, duration, conflicts, reason = generate_schedule(
                    phase=phase,
                    progress_callback=progress_cb, 
                    stop_check=stop_check
                )
                
                # Treat generation as success even with minor conflicts,
                # as long as it wasn't a crash or user cancellation.
                # User cancellation is success=False.
                actual_success = success if not stop_check() else False
                if not actual_success and not stop_check() and conflicts < 20:
                    actual_success = True # Partial Success
                
                if actual_success:
                    # Update the official generation timestamp
                    now_str = datetime.now().strftime('%B %d, %Y %I:%M %p')
                    ls_time = Setting.query.filter_by(key='last_setup_time').first()
                    if ls_time: ls_time.value = now_str
                    else: db.session.add(Setting(key='last_setup_time', value=now_str))
                    db.session.commit()

                scheduler_status['last_result'] = {
                    'success': actual_success,
                    'duration': duration,
                    'conflicts': conflicts,
                    'reason': reason,
                    'is_partial': (conflicts > 0 and not stop_check()),
                    'is_cancelled': stop_check()
                }
            except Exception as e:
                scheduler_status['message'] = f"Error: {str(e)}"
                scheduler_status['last_result'] = {'success': False, 'reason': str(e)}
            finally:
                scheduler_status['running'] = False
                scheduler_status['progress'] = 100

    # Start thread with app context
    thread = threading.Thread(target=run_wrapper, args=(app.app_context(), phase))
    thread.start()
    
    return jsonify({'success': True, 'msg': f'Generation for {phase.upper()} started.'})

@app.route('/admin/teacher/toggle_hybrid/<int:id>', methods=['POST'])
@login_required
def admin_teacher_toggle_hybrid(id):
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    teacher = Teacher.query.get_or_404(id)
    teacher.is_hybrid = not teacher.is_hybrid
    db.session.commit()
    return jsonify({
        'success': True, 
        'is_hybrid': teacher.is_hybrid,
        'msg': f"Teacher {teacher.name} is now {'Hybrid' if teacher.is_hybrid else 'Standard'}."
    })

@app.route('/admin/schedule/cancel', methods=['POST'])
@login_required
def admin_cancel_schedule():
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    global scheduler_status
    if scheduler_status['running']:
        scheduler_status['cancel_requested'] = True
        scheduler_status['message'] = 'Stopping engine...'
        return jsonify({'success': True, 'msg': 'Cancellation requested.'})
    return jsonify({'error': 'No generation in progress.'}), 400

@app.route('/admin/schedule/status')
@login_required
def admin_schedule_status():
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    return jsonify(scheduler_status)

@app.route('/admin/setup', methods=['GET'])
@login_required
def admin_setup():
    if current_user.role != 'admin': return redirect(url_for('index'))
    
    settings = {s.key: s.value for s in Setting.query.all()}
    # Consistent sorting for the generation view
    sections = Section.query.all()
    sections.sort(key=lambda x: (x.department, natural_sort_key(str(x.grade_level)), natural_sort_key(x.name)))
    teachers = Teacher.query.all()
    teachers.sort(key=lambda x: (x.department, natural_sort_key(x.name)))
    classrooms = Classroom.query.all()
    classrooms.sort(key=lambda x: (x.building, natural_sort_key(x.name)))
    
    # Pre-calculate unique grade levels for grouping in templates
    jhs_grades = sorted(list(set(s.grade_level for s in sections if s.department == 'JHS')), key=natural_sort_key)
    shs_grades = sorted(list(set(s.grade_level for s in sections if s.department == 'SHS')), key=natural_sort_key)
    
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    schedules = Schedule.query.filter_by(run_id=active_run.id).all() if active_run else []
    runs = ScheduleRun.query.order_by(ScheduleRun.id.desc()).all()
    
    # Global bounds for consistent timetable display
    g_start, g_end = get_global_schedule_bounds()
    
    section_grids = {s.id: prepare_schedule_grid(s.id, 'section', schedules, force_start=g_start, force_end=g_end) for s in sections}
    teacher_grids = {t.id: prepare_schedule_grid(t.id, 'teacher', schedules, force_start=g_start, force_end=g_end) for t in teachers}
    room_grids = {r.id: prepare_schedule_grid(r.id, 'classroom', schedules, force_start=g_start, force_end=g_end) for r in classrooms}

    return render_template('admin_setup.html', settings=settings, 
                           sections=sections, teachers=teachers, classrooms=classrooms,
                           jhs_grades=jhs_grades, shs_grades=shs_grades,
                           section_grids=section_grids, teacher_grids=teacher_grids, room_grids=room_grids,
                           runs=runs, active_run=active_run)

@app.route('/admin/schedule/restore/<int:run_id>', methods=['POST'])
@login_required
def restore_schedule(run_id):
    if current_user.role != 'admin': return redirect(url_for('index'))
    run_to_restore = ScheduleRun.query.get_or_404(run_id)
    
    runs = ScheduleRun.query.all()
    for run in runs:
        run.is_active = False
        
    run_to_restore.is_active = True
    db.session.commit()
    
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'success': True, 'msg': 'Schedule restored successfully!'})
        
    flash('Previous schedule restored and set as active.')
    return redirect(url_for('admin_setup'))

@app.route('/teacher')
@login_required
def teacher_dashboard():
    if current_user.role != 'teacher': return redirect(url_for('index'))
    teacher = None
    if getattr(current_user, 'related_id', None):
        teacher = Teacher.query.get(current_user.related_id)
        
    grid_data = None
    advisees = []
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    if teacher and active_run:
        try:
            schedules = Schedule.query.filter_by(teacher_id=teacher.id, run_id=active_run.id).all()
            g_start, g_end = get_global_schedule_bounds()
            grid_data = prepare_schedule_grid(teacher.id, 'teacher', schedules, force_start=g_start, force_end=g_end)
            advisees = Section.query.filter_by(adviser_id=teacher.id).all()
        except Exception as e:
            print(f"Error preparing teacher grid: {e}")
            grid_data = None
        
    teacher_load = 0
    school_days = 5
    if teacher:
        active_days_setting = Setting.query.filter_by(key='active_days').first()
        if active_days_setting:
            school_days = len(active_days_setting.value.split(','))
        
        if active_run:
            schedules = Schedule.query.filter_by(teacher_id=teacher.id, run_id=active_run.id).all()
            for sch in schedules:
                teacher_load += sch.subject.duration_mins
            teacher_load = round(teacher_load / 60, 1)

    last_setup_time = active_run.created_at.strftime('%B %d, %Y %I:%M %p') if active_run else "Not yet generated"

    return render_template('teacher_dashboard.html', teacher=teacher, grid_data=grid_data, advisees=advisees, last_setup_time=last_setup_time, teacher_load=teacher_load, school_days=school_days)

@app.route('/student')
@login_required
def student_dashboard():
    if current_user.role != 'student': return redirect(url_for('index'))
    section = None
    if getattr(current_user, 'related_id', None):
        section = Section.query.get(current_user.related_id)
        
    shift_info = "Whole Day"
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    if section and active_run:
        schedules = Schedule.query.filter_by(section_id=section.id, run_id=active_run.id).all()
        g_start, g_end = get_global_schedule_bounds()
        grid_data = prepare_schedule_grid(section.id, 'section', schedules, force_start=g_start, force_end=g_end)
        
        # Detect shift
        if section.department == 'JHS':
            settings = {s.key: s.value for s in Setting.query.all()}
            if settings.get(f'jhs_am_grade_{section.grade_level}') in ['active', 'on']:
                shift_info = "AM Shift"
            else:
                shift_info = "PM Shift"
        else:
            shift_info = "Whole Day"
            
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    last_setup_time = active_run.created_at.strftime('%B %d, %Y %I:%M %p') if active_run else "Not yet generated"
    return render_template('student_dashboard.html', section=section, grid_data=grid_data, shift_info=shift_info, last_setup_time=last_setup_time)

@app.route('/export_pdf/<view_type>/<int:id>')
@login_required
def export_pdf(view_type, id):
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    schedules = []
    run_filter = active_run.id if active_run else -1
    if view_type == 'teacher':
        entity = Teacher.query.get_or_404(id)
        schedules = Schedule.query.filter_by(teacher_id=id, run_id=run_filter).all()
    elif view_type == 'section':
        entity = Section.query.get_or_404(id)
        schedules = Schedule.query.filter_by(section_id=id, run_id=run_filter).all()
    elif view_type == 'classroom':
        entity = Classroom.query.get_or_404(id)
        schedules = Schedule.query.filter_by(room_id=id, run_run=run_filter).all() if False else Schedule.query.filter_by(room_id=id, run_id=run_filter).all()
    else:
        return "Invalid type", 400
        
    if view_type == 'classroom':
        settings_dict = {s.key: s.value for s in Setting.query.all()}
        am_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'AM']
        pm_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'PM']
        
        grids_data = []
        if am_sch and pm_sch:
            grids_data.append({'grid': prepare_condensed_grid(id, view_type, am_sch), 'shift': 'AM Schedule'})
            grids_data.append({'grid': prepare_condensed_grid(id, view_type, pm_sch), 'shift': 'PM Schedule'})
        else:
            grids_data.append({'grid': prepare_condensed_grid(id, view_type, schedules), 'shift': ''})
    else:
        grids_data = [{'grid': prepare_condensed_grid(id, view_type, schedules), 'shift': ''}]

    template = 'schedule_print_condensed.html'

    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    last_setup_time = active_run.created_at.strftime('%B %d, %Y %I:%M %p') if active_run else "Not yet generated"
    school_name = Setting.query.filter_by(key='school_name').first()
    school_name = school_name.value if school_name else "Andres M. Luciano High School"
    school_year = Setting.query.filter_by(key='school_year').first()
    school_year = school_year.value if school_year else "2023-2024"

    return render_template(template, entity=entity, grids_data=grids_data, 
                           view_type=view_type, last_setup_time=last_setup_time,
                           school_name=school_name, school_year=school_year)

@app.route('/export_pdf_bulk/<view_type>/<filter_value>', strict_slashes=False)
@login_required
def export_pdf_bulk(view_type, filter_value):
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    run_filter = active_run.id if active_run else -1
    schedules = Schedule.query.filter_by(run_id=run_filter).all()
    
    entities = []
    if view_type == 'section':
        if filter_value in ['JHS', 'SHS']: entities = Section.query.filter_by(department=filter_value).all()
        else: entities = Section.query.all()
    elif view_type == 'teacher':
        if filter_value == 'All':
            entities = Teacher.query.filter(Teacher.name != 'TBA').order_by(Teacher.name).all()
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
        else:
            active_ids = {sch.teacher_id for sch in schedules if sch.section.department == filter_value}
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
    elif view_type == 'classroom':
        if filter_value == 'All':
            entities = Classroom.query.all()
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.room_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.room_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()
        else:
            active_ids = {sch.room_id for sch in schedules if sch.section.department == filter_value}
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()

    if not active_run or not entities:
        flash("No active schedule or matching entities found.", "error")
        return redirect(url_for('admin_setup'))

    # Prepare all grids and attach metadata
    # NOTE: No global forcing here as per user request for condensed individual exports
    bulk_data = []
    for entity in entities:
        is_shs = True # Force condensed format for all bulk exports
        if view_type == 'teacher':
            ent_sch = [s for s in schedules if s.teacher_id == entity.id]
            bulk_data.append({'entity': entity, 'grid_data': prepare_condensed_grid(entity.id, view_type, ent_sch), 'is_condensed': True})
        elif view_type == 'section':
            ent_sch = [s for s in schedules if s.section_id == entity.id]
            bulk_data.append({'entity': entity, 'grid_data': prepare_condensed_grid(entity.id, view_type, ent_sch), 'is_condensed': True})
        else:
            ent_sch = [s for s in schedules if s.room_id == entity.id]
            settings_dict = {s.key: s.value for s in Setting.query.all()}
            am_sch = [s for s in ent_sch if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'AM']
            pm_sch = [s for s in ent_sch if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'PM']
            
            if am_sch and pm_sch:
                bulk_data.append({'entity': entity, 'grid_data': prepare_condensed_grid(entity.id, view_type, am_sch), 'is_condensed': True, 'shift_label': 'AM Schedule'})
                bulk_data.append({'entity': entity, 'grid_data': prepare_condensed_grid(entity.id, view_type, pm_sch), 'is_condensed': True, 'shift_label': 'PM Schedule'})
            else:
                bulk_data.append({'entity': entity, 'grid_data': prepare_condensed_grid(entity.id, view_type, ent_sch), 'is_condensed': True})

    school_name = Setting.query.filter_by(key='school_name').first()
    school_name = school_name.value if school_name else "Andres M. Luciano High School"
    school_year = Setting.query.filter_by(key='school_year').first()
    school_year = school_year.value if school_year else "2023-2024"
    last_setup_time = active_run.created_at.strftime('%B %d, %Y %I:%M %p')

    return render_template('schedule_print_bulk.html', 
                           bulk_data=bulk_data, view_type=view_type,
                           school_name=school_name, school_year=school_year,
                           last_setup_time=last_setup_time)

def write_schedule_to_excel(writer, entity, view_type, schedules):
    grid_data = prepare_condensed_grid(entity.id, view_type, schedules)

    days = grid_data['days']
    slots = grid_data['time_slots']
    
    grids_to_write = []
    if view_type == 'classroom':
        settings_dict = {s.key: s.value for s in Setting.query.all()}
        am_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'AM']
        pm_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'PM']
        if am_sch and pm_sch:
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, am_sch), f"{entity.name[:18]} - AM"))
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, pm_sch), f"{entity.name[:18]} - PM"))
        else:
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, schedules), entity.name[:31]))
    else:
        grids_to_write.append((prepare_condensed_grid(entity.id, view_type, schedules), entity.name[:31]))

    for grid_data, sheet_name in grids_to_write:
        days = grid_data['days']
        slots = grid_data['time_slots']
        occupied = grid_data['occupied']
        spans = grid_data['spans']
        grid = grid_data['grid']

        safe_sheet_name = sheet_name.replace('/', '-').replace('\\', '-').replace('?', '').replace('*', '').replace('[', '').replace(']', '')
        if not safe_sheet_name: safe_sheet_name = "Sheet"
        
        worksheet = writer.book.add_worksheet(safe_sheet_name)
        worksheet.hide_gridlines(2)

        header_format = writer.book.add_format({
            'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1, 'bg_color': '#D3D3D3', 'text_wrap': True
        })
        cell_format = writer.book.add_format({
            'align': 'center', 'valign': 'vcenter', 'border': 1, 'text_wrap': True
        })
        time_format = writer.book.add_format({
            'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1, 'bg_color': '#F0F0F0'
        })
        break_format = writer.book.add_format({
            'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1, 'bg_color': '#FFFFE0'
        })

        worksheet.write(0, 0, "Time Interval", header_format)
        for c, day in enumerate(days, 1):
            worksheet.write(0, c, day, header_format)

        for r, slot in enumerate(slots, 1):
            worksheet.write(r, 0, f"{slot} - {grid_data['slot_ranges'][slot]}", time_format)
            for c, day in enumerate(days, 1):
                if not occupied[slot][day]:
                    sch = grid[slot][day]
                    if sch in ['BREAK', 'LUNCH', 'VACANT']:
                        span = spans[slot][day]
                        if span > 1:
                            worksheet.merge_range(r, c, r + span - 1, c, sch, break_format)
                        else:
                            worksheet.write(r, c, sch, break_format)
                    elif sch:
                        if view_type == 'teacher':
                            text = f"{sch.subject.name}\nGrade {sch.section.grade_level} - {sch.section.name}\nRM: {sch.room.name}"
                        elif view_type == 'section':
                            text = f"{sch.subject.name}\n{sch.teacher.name}\n{sch.room.name}"
                        else: # classroom
                            text = f"{sch.subject.name}\nSection {sch.section.name}\n{sch.teacher.name}"
                        span = spans[slot][day]
                        if span > 1:
                            worksheet.merge_range(r, c, r + span - 1, c, text, cell_format)
                        else:
                            worksheet.write(r, c, text, cell_format)
                    else:
                        worksheet.write(r, c, "", cell_format)

        worksheet.set_column('A:A', 18)
        for col_idx in range(1, len(days) + 1):
            col_letter = chr(65 + col_idx)
            worksheet.set_column(f'{col_letter}:{col_letter}', 25)

def write_schedule_to_word(doc, entity, view_type, schedules, school_name, school_year):
    grids_to_write = []
    if view_type == 'classroom':
        settings_dict = {s.key: s.value for s in Setting.query.all()}
        am_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'AM']
        pm_sch = [s for s in schedules if get_section_shift(s.section.department, s.section.grade_level, settings_dict) == 'PM']
        if am_sch and pm_sch:
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, am_sch), f"AM Schedule"))
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, pm_sch), f"PM Schedule"))
        else:
            grids_to_write.append((prepare_condensed_grid(entity.id, view_type, schedules), ""))
    else:
        grids_to_write.append((prepare_condensed_grid(entity.id, view_type, schedules), ""))

    for i, (grid_data, shift_label) in enumerate(grids_to_write):
        if i > 0:
            doc.add_page_break()

        days = grid_data['days']
        slots = grid_data['time_slots']

        # Header
        header = doc.add_heading(school_name, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Academic Year: {school_year}")
        run.italic = True
        
        ent_title = entity.name
        if shift_label:
            ent_title += f" – {shift_label}"

        title = doc.add_heading(ent_title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        details_p = doc.add_paragraph()
        details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if view_type == 'section':
            adviser_name = entity.adviser.name if entity.adviser else 'TBA'
            room_name = entity.room.name if entity.room else 'TBA'
            details_text = f"SECTION: {entity.name}  |  GRADE: {entity.grade_level}"
            if entity.track:
                details_text += f"  |  TRACK: {entity.track}"
            details_text += f"  |  ADVISER: {adviser_name}  |  ROOM: {room_name}"
        elif view_type == 'teacher':
            details_text = f"TEACHER: {entity.name}  |  DEPT: {entity.department}"
        elif view_type == 'classroom':
            details_text = f"ROOM: {entity.name}  |  BUILDING: {entity.building}"
        else:
            details_text = ""
            
        if details_text:
            details_run = details_p.add_run(details_text)
            details_run.bold = True
            
        doc.add_paragraph() # Spacer

        # Create Table
        table = doc.add_table(rows=len(slots) + 1, cols=len(days) + 1)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Time'
        for j, day in enumerate(days):
            hdr_cells[j+1].text = day
            
        # Styling headers
        def set_cell_background(cell, fill_color):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fill_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for cell in hdr_cells:
            set_cell_background(cell, "1F4E78")
            if cell.paragraphs and cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fill Data
        for s_idx, slot in enumerate(slots):
            row_cells = table.rows[s_idx+1].cells
            row_cells[0].text = slot
            for d_idx, day in enumerate(days):
                if not grid_data['occupied'][slot][day]:
                    cell = row_cells[d_idx+1]
                    sch = grid_data['grid'][slot][day]
                    span = grid_data['spans'][slot][day]
                    
                    if isinstance(sch, str):
                        cell.text = sch
                        set_cell_background(cell, "F1F5FA")
                    elif sch:
                        if view_type == 'teacher':
                            cell.text = f"{sch.subject.name}\nGrade {sch.section.grade_level} - {sch.section.name}\nRM: {sch.room.name}"
                        elif view_type == 'section':
                            cell.text = f"{sch.subject.name}\n{sch.teacher.name}\n{sch.room.name}"
                        else: # classroom
                            cell.text = f"{sch.subject.name}\n{sch.teacher.name}\nSection {sch.section.name}"
                        set_cell_background(cell, "DCFCE7")
                    
                    # Internal Alignment
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

                    if span > 1:
                        # In python-docx, we merge cells manually
                        for k in range(1, span):
                            if (s_idx + 1 + k) < len(table.rows):
                                cell.merge(table.rows[s_idx+1+k].cells[d_idx+1])

@app.route('/export_word/<view_type>/<int:id>')
@login_required
def export_word(view_type, id):
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    schedules = []
    run_filter = active_run.id if active_run else -1
    if view_type == 'teacher':
        entity = Teacher.query.get_or_404(id)
        schedules = Schedule.query.filter_by(teacher_id=id, run_id=run_filter).all()
    elif view_type == 'section':
        entity = Section.query.get_or_404(id)
        schedules = Schedule.query.filter_by(section_id=id, run_id=run_filter).all()
    elif view_type == 'classroom':
        entity = Classroom.query.get_or_404(id)
        schedules = Schedule.query.filter_by(room_id=id, run_id=run_filter).all()
    else:
        return "Invalid type", 400
        
    school_name = Setting.query.filter_by(key='school_name').first()
    school_name = school_name.value if school_name else "Andres M. Luciano High School"
    school_year = Setting.query.filter_by(key='school_year').first()
    school_year = school_year.value if school_year else "2023-2024"

    doc = docx.Document()
    write_schedule_to_word(doc, entity, view_type, schedules, school_name, school_year)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    safe_entity_name = sanitize_filename(entity.name)
    filename = f"Schedule_{safe_entity_name}.docx"
    return send_file(output, as_attachment=True, download_name=filename)

@app.route('/export_word_bulk/<view_type>/<filter_value>', strict_slashes=False)
@login_required
def export_word_bulk(view_type, filter_value):
    if current_user.role != 'admin': return redirect(url_for('index'))
    
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    run_filter = active_run.id if active_run else -1
    schedules = Schedule.query.filter_by(run_id=run_filter).all()
    
    entities = []
    if view_type == 'section':
        if filter_value == 'All': entities = Section.query.all()
        else: entities = Section.query.filter_by(department=filter_value).all()
    elif view_type == 'teacher':
        if filter_value == 'All':
            entities = [t for t in Teacher.query.all() if t.name != 'TBA']
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
        else:
            active_ids = {sch.teacher_id for sch in schedules if sch.section.department == filter_value}
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
    elif view_type == 'classroom':
        if filter_value == 'All':
            entities = Classroom.query.all()
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.room_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.room_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()
        else:
            active_ids = {sch.room_id for sch in schedules if sch.section.department == filter_value}
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()
            
    if not active_run:
        flash("No active schedule found. Please generate or restore a schedule first.", "error")
        return redirect(url_for('admin_setup'))

    if not schedules:
        flash("The active schedule run contains no assignments.", "error")
        return redirect(url_for('admin_setup'))

    if not entities:
        flash("No entities found for the selected filter.", "error")
        return redirect(url_for('admin_setup'))
        
    school_name = Setting.query.filter_by(key='school_name').first()
    school_name = school_name.value if school_name else "Andres M. Luciano High School"
    school_year = Setting.query.filter_by(key='school_year').first()
    school_year = school_year.value if school_year else "2023-2024"

    doc = docx.Document()
    for i, entity in enumerate(entities):
        write_schedule_to_word(doc, entity, view_type, schedules, school_name, school_year)
        if i < len(entities) - 1:
            doc.add_page_break()
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    safe_filter = sanitize_filename(filter_value)
    safe_view = sanitize_filename(view_type)
    filename = f"Bulk_Schedule_{safe_filter}_{safe_view}.docx"
    return send_file(output, as_attachment=True, download_name=filename)

@app.route('/export_excel/<view_type>/<int:id>')
@login_required
def export_excel(view_type, id):
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    schedules = []
    run_filter = active_run.id if active_run else -1
    if view_type == 'teacher':
        entity = Teacher.query.get_or_404(id)
        schedules = Schedule.query.filter_by(teacher_id=id, run_id=run_filter).all()
    elif view_type == 'section':
        entity = Section.query.get_or_404(id)
        schedules = Schedule.query.filter_by(section_id=id, run_id=run_filter).all()
    elif view_type == 'classroom':
        entity = Classroom.query.get_or_404(id)
        schedules = Schedule.query.filter_by(room_id=id, run_id=run_filter).all()
    else:
        return "Invalid type", 400
        
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        write_schedule_to_excel(writer, entity, view_type, schedules)
            
    output.seek(0)
    safe_name = sanitize_filename(entity.name)
    filename = f"Schedule_{safe_name}.xlsx"
    return send_file(output, as_attachment=True, download_name=filename)

@app.route('/export_excel_bulk/<view_type>/<filter_value>', strict_slashes=False)
@login_required
def export_excel_bulk(view_type, filter_value):
    if current_user.role != 'admin': return redirect(url_for('index'))
    
    active_run = ScheduleRun.query.filter_by(is_active=True).first()
    run_filter = active_run.id if active_run else -1
    schedules = Schedule.query.filter_by(run_id=run_filter).all()
    
    entities = []
    if view_type == 'section':
        if filter_value == 'All':
            entities = Section.query.all()
        else:
            entities = Section.query.filter_by(department=filter_value).all()
    elif view_type == 'teacher':
        if filter_value == 'All':
            entities = [t for t in Teacher.query.all() if t.name != 'TBA']
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.teacher_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
        else:
            active_ids = {sch.teacher_id for sch in schedules if sch.section.department == filter_value}
            entities = [t for t in Teacher.query.filter(Teacher.id.in_(active_ids)).all() if t.name != 'TBA']
    elif view_type == 'classroom':
        if filter_value == 'All':
            entities = Classroom.query.all()
        elif filter_value == 'Both':
            # Strict Dual-Presence (Intersection of JHS and SHS)
            jhs_ids = {sch.room_id for sch in schedules if sch.section.department == 'JHS'}
            shs_ids = {sch.room_id for sch in schedules if sch.section.department == 'SHS'}
            active_ids = jhs_ids.intersection(shs_ids)
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()
        else:
            active_ids = {sch.room_id for sch in schedules if sch.section.department == filter_value}
            entities = Classroom.query.filter(Classroom.id.in_(active_ids)).all()
            
    if not active_run:
        flash("No active schedule found. Please generate or restore a schedule first.", "error")
        return redirect(url_for('admin_setup'))

    if not schedules:
        flash("The active schedule run contains no assignments.", "error")
        return redirect(url_for('admin_setup'))

    if not entities:
        flash("No entities found for the selected filter.", "error")
        return redirect(url_for('admin_setup'))
        
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for entity in entities:
            # Sheet names must be unique. If names collide, add ID.
            write_schedule_to_excel(writer, entity, view_type, schedules)
            
    output.seek(0)
    safe_filter = sanitize_filename(filter_value)
    safe_view = sanitize_filename(view_type)
    filename = f"Bulk_Schedule_{safe_filter}_{safe_view}.xlsx"
    return send_file(output, as_attachment=True, download_name=filename)

@app.route('/admin/teacher/toggle_status/<int:id>', methods=['POST'])
@login_required
def admin_toggle_teacher_status(id):
    if current_user.role != 'admin': return jsonify({'error': 'Unauthorized'}), 403
    teacher = Teacher.query.get_or_404(id)
    teacher.is_active = not teacher.is_active
    db.session.commit()
    return jsonify({
        'success': True, 
        'is_active': teacher.is_active, 
        'msg': f"Teacher profile {'activated' if teacher.is_active else 'disabled'} successfully."
    })

 # Initialize DB command
@app.cli.command('initdb')
def initdb_command():
    """Creates the database tables and admin user."""
    db.create_all()
    # Create default admin user
    if not User.query.filter_by(username='admin').first():
        admin = User(
            username='admin',
            password=generate_password_hash('admin123'),
            role='admin'
        )
        db.session.add(admin)
        db.session.commit()
    print('Initialized the database.')

@app.route('/admin/import/template/<module>')
@login_required
def download_import_template(module):
    if current_user.role != 'admin':
        return redirect(url_for('index'))
        
    templates = {
        'teachers': {
            'headers': ['Name', 'Department', 'Is Master', 'Max Hours', 'Stay Window', 'Handle Sec A', 'Preferred Days', 'Subjects', 'Grade Levels'],
            'sample': ['Sample Teacher', 'JHS', 'No', '6', '9', 'No', 'M,T,W,Th,F', 'Math, Science', '7,8']
        },
        'classrooms': {
            'headers': ['Room Name', 'Building', 'Type'],
            'sample': ['Sample Room', 'Main', 'Lecture']
        },
        'sections': {
            'headers': ['Section Name', 'Department', 'Grade Level', 'Track', 'Adviser', 'Room', 'Is Section A'],
            'sample': ['Sample Section', 'JHS', '7', '', 'Sample Teacher', 'Sample Room', 'No']
        },
        'subjects': {
            'headers': ['Subject Name', 'Department', 'Duration Mins', 'Meetings Per Week', 'Requires Lab', 'Grade Level', 'Track'],
            'sample': ['Sample Subject', 'JHS', '60', '4', 'No', '7', '']
        }
    }
    
    if module not in templates:
        flash('Invalid template requested.', 'error')
        return redirect(url_for('admin_dashboard'))
        
    data = templates[module]
    df = pd.DataFrame([data['sample']], columns=data['headers'])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Template')
    
    output.seek(0)
    
    return send_file(
        output,
        download_name=f"{module}_template.xlsx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=False)
